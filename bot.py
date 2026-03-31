import pandas as pd
import logging
import json
import os
import sys
import zipfile
import shutil
import asyncio
import uuid
from datetime import datetime
from pathlib import Path
from io import BytesIO

try:
    from fastapi import FastAPI, File, UploadFile, Form, HTTPException
    from fastapi.middleware.cors import CORSMiddleware
    from fastapi.responses import FileResponse
    from starlette.background import BackgroundTask
    import uvicorn
    FASTAPI_AVAILABLE = True
except ImportError:
    FASTAPI_AVAILABLE = False

from docx import Document
from docx.shared import Pt
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, WebAppInfo, ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    CallbackQueryHandler,
    ConversationHandler,
    filters,
    ContextTypes,
)

# ---------------------------------------------------------------------------
# Логирование
# ---------------------------------------------------------------------------
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Конфигурация
# ---------------------------------------------------------------------------
BOT_TOKEN = "8603254098:AAFvj8bxbbRrEf_8emQh4JlSO9rHELcujJI"

ADMIN_ID = 630597358  # Главный администратор
import os
API_PORT = int(os.getenv("PORT", 8080))

# Состояния ConversationHandler
(
    CHOOSING,
    WAITING_FILES,
    WAITING_SERIALS,
    CONFIRM_DATA,
    WAITING_TEMPLATE,
    WAITING_BARCODE_PHOTO,
    ADMIN_PANEL,
    ADMIN_ADD_USER,
    WAITING_QR_INPUT,           # ожидание ввода серийных номеров для QR
    SUPPLY_WAITING_PHOTO,       # «Новая поставка»: ожидание фото коробки
    SUPPLY_WAITING_NUMBER,      # «Новая поставка»: ожидание ввода номера коробки
    SUPPLY_EDITING_SERIALS,     # «Новая поставка»: ручное редактирование серийников
) = range(12)

# Пути
BASE_DIR = Path(__file__).parent
TEMP_DIR = BASE_DIR / "temp_files"
TEMPLATES_DIR = BASE_DIR / "templates"
USER_TEMPLATES_DIR = TEMPLATES_DIR / "user_templates"
DEFAULT_TEMPLATE_PATH = TEMPLATES_DIR / "акт_демонтажа_шаблон.docx"
USERS_FILE = BASE_DIR / "allowed_users.json"

for d in (TEMP_DIR, TEMPLATES_DIR, USER_TEMPLATES_DIR):
    d.mkdir(exist_ok=True)

# Хранилище сессий
user_sessions: dict = {}


# ===========================================================================
#  УПРАВЛЕНИЕ ПОЛЬЗОВАТЕЛЯМИ (JSON-файл)
# ===========================================================================

def load_allowed_users() -> dict:
    """Загружает список разрешённых пользователей из файла."""
    if USERS_FILE.exists():
        try:
            with open(USERS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def save_allowed_users(users: dict):
    """Сохраняет список разрешённых пользователей в файл."""
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=2)


def is_allowed(user_id: int) -> bool:
    """Проверяет, есть ли у пользователя доступ к боту."""
    if user_id == ADMIN_ID:
        return True
    users = load_allowed_users()
    return str(user_id) in users


def add_user(user_id: int, name: str = ""):
    """Добавляет пользователя в список разрешённых."""
    users = load_allowed_users()
    users[str(user_id)] = {"name": name, "added": datetime.now().strftime("%d.%m.%Y %H:%M")}
    save_allowed_users(users)


def remove_user(user_id: int):
    """Удаляет пользователя из списка разрешённых."""
    users = load_allowed_users()
    users.pop(str(user_id), None)
    save_allowed_users(users)


def get_all_users() -> dict:
    """Возвращает всех разрешённых пользователей."""
    return load_allowed_users()


# ===========================================================================
#  СКАНИРОВАНИЕ ШТРИХКОДОВ
# ===========================================================================

def _is_numeric_serial(value: str) -> bool:
    """Серийный номер — строка ТОЛЬКО из цифр длиной >= 6."""
    v = value.strip()
    return v.isdigit() and len(v) >= 6


def _get_tiles(img):
    """
    Специализированная нарезка для фото коробок Нартис (2 столбца × 9 строк наклеек).
    Стратегия:
      A) Обзорные тайлы (полная ширина): полное фото + 4 полосы + 3 сдвинутых полосы
      B) Колоночные тайлы: левый/правый столбец × 9 полос с перекрытием 15%
      C) Апскейл каждого тайла до min-ширины 900px для лучшей читаемости штрихкода
    Итого ~30 тайлов — максимальное покрытие всех 18 наклеек.
    """
    from PIL import Image as PILImage
    w, h = img.size
    ov_x = int(w * 0.10)
    ov_y = int(h * 0.10)

    added = set()
    tiles = []

    def add(t):
        key = (t.size, t.tobytes()[:128])
        if key in added:
            return
        added.add(key)
        tiles.append(t)
        # Апскейл: если тайл уже 900px — масштабируем для лучшей читаемости штрихкода
        tw, th = t.size
        if tw < 900:
            scale = 900 / tw
            scaled = t.resize((int(tw * scale), int(th * scale)), PILImage.LANCZOS)
            tiles.append(scaled)

    # ── A1. Полное изображение ─────────────────────────────────────────────────
    tiles.append(img)

    # ── A2. 4 горизонтальные полосы (основная сетка) ──────────────────────────
    qh = h // 4
    for i in range(4):
        y0 = max(0, i * qh - ov_y)
        y1 = min(h, (i + 1) * qh + ov_y)
        add(img.crop((0, y0, w, y1)))

    # ── A3. 3 горизонтальные полосы (сдвинутая сетка — ловит стыки A2) ────────
    th3 = h // 3
    for i in range(3):
        y0 = max(0, i * th3 - ov_y)
        y1 = min(h, (i + 1) * th3 + ov_y)
        add(img.crop((0, y0, w, y1)))

    # ── B. Два столбца × 9 рядов (по одному на каждую наклейку) ─────────────
    ROWS = 9
    col_l = img.crop((0, 0, w // 2 + ov_x, h))
    col_r = img.crop((w // 2 - ov_x, 0, w, h))

    for col in (col_l, col_r):
        cw, ch = col.size
        row_h = ch // ROWS
        row_ov = int(ch * 0.15)
        for i in range(ROWS):
            y0 = max(0, i * row_h - row_ov)
            y1 = min(ch, (i + 1) * row_h + row_ov)
            add(col.crop((0, y0, cw, y1)))

    return tiles


def _preprocess_variants(img):
    """
    Набор вариантов предобработки одного тайла.
    Каждый вариант — PIL Image в режиме RGB.
    """
    from PIL import ImageEnhance, ImageFilter, ImageOps
    rgb = img.convert("RGB")
    gray = img.convert("L")
    variants = []

    # 1. Оригинал
    variants.append(rgb)

    # 2. Autocontrast — автоматически растягивает гистограмму
    variants.append(ImageOps.autocontrast(gray, cutoff=1).convert("RGB"))

    # 3. Высокий контраст + двойная резкость
    high = ImageEnhance.Contrast(gray).enhance(3.0)
    high = ImageEnhance.Sharpness(high.convert("RGB")).enhance(2.5)
    variants.append(high)

    # 4. Unsharp Mask — лучший способ сделать штрихкоды чёткими без артефактов
    usm = rgb.filter(ImageFilter.UnsharpMask(radius=2, percent=200, threshold=2))
    variants.append(usm)

    # 5. Умеренный контраст + резкость (мягкий вариант)
    mild = ImageEnhance.Contrast(gray).enhance(1.8).filter(ImageFilter.SHARPEN)
    variants.append(mild.convert("RGB"))

    return variants


def _pyzbar_scan_tile(img) -> list:
    """Сканирует один тайл через pyzbar со всеми вариантами предобработки."""
    from pyzbar.pyzbar import decode as pyz_decode
    results = []
    seen = set()
    for variant in _preprocess_variants(img):
        for b in pyz_decode(variant):
            try:
                val = b.data.decode("utf-8").strip()
            except Exception:
                val = b.data.decode("latin-1", errors="replace").strip()
            if val and val not in seen:
                seen.add(val)
                results.append((val, b.rect.top))
    return results


def _zxingcpp_scan_tile(img) -> list:
    """Сканирует один тайл через zxingcpp со всеми вариантами предобработки."""
    import zxingcpp
    import numpy as np
    results = []
    seen = set()
    for variant in _preprocess_variants(img):
        arr = np.array(variant)
        for b in zxingcpp.read_barcodes(arr):
            val = b.text.strip()
            if val and val not in seen:
                seen.add(val)
                results.append((val, b.position.top_left.y))
    return results


def scan_barcodes(image_path: str) -> dict:
    """
    Максимально полное сканирование штрихкодов:
    - ~21 тайл с перекрытием по сетке 2-столбца × 3-4 строки
    - 5 вариантов предобработки на каждый тайл
    - Обе библиотеки параллельно (pyzbar + zxingcpp)
    Специально оптимизировано для фото коробок со счётчиками Нартис.
    """
    result = {"ok": False, "method": "", "all": [], "serials": [], "skipped": 0, "error": ""}
    all_pairs: list = []
    methods_used: list = []

    try:
        from PIL import Image as PILImage
        img = PILImage.open(image_path)
        # Корректируем EXIF-ориентацию (iPhone часто даёт повёрнутые фото)
        try:
            from PIL import ImageOps
            img = ImageOps.exif_transpose(img)
        except Exception:
            pass
        img = img.convert("RGB")
        tiles = _get_tiles(img)
    except Exception as e:
        result["error"] = str(e)
        return result

    # ── pyzbar ──────────────────────────────────────────────────────────────
    pyzbar_available = False
    try:
        from pyzbar.pyzbar import decode as _pyz_check  # noqa
        pyzbar_available = True
    except ImportError:
        pass

    if pyzbar_available:
        try:
            for tile in tiles:
                all_pairs.extend(_pyzbar_scan_tile(tile))
            methods_used.append("pyzbar")
        except Exception as e:
            logger.error(f"pyzbar error: {e}")

    # ── zxingcpp ─────────────────────────────────────────────────────────────
    zxingcpp_available = False
    try:
        import zxingcpp as _zx_check  # noqa
        zxingcpp_available = True
    except ImportError:
        pass

    if zxingcpp_available:
        try:
            for tile in tiles:
                all_pairs.extend(_zxingcpp_scan_tile(tile))
            methods_used.append("zxingcpp")
        except Exception as e:
            logger.error(f"zxingcpp error: {e}")

    if not pyzbar_available and not zxingcpp_available:
        result["error"] = "no_libs"
        return result

    result["ok"] = True
    result["method"] = "+".join(methods_used) if methods_used else "unknown"

    # Дедупликация: приоритет записи с наименьшей Y-координатой (сверху вниз)
    seen_vals: set = set()
    unique_pairs: list = []
    for val, y in sorted(all_pairs, key=lambda x: x[1]):
        if val not in seen_vals:
            seen_vals.add(val)
            unique_pairs.append((val, y))

    result["all"] = [v for v, _ in unique_pairs]

    seen_serials: set = set()
    for val in result["all"]:
        if _is_numeric_serial(val) and val not in seen_serials:
            seen_serials.add(val)
            result["serials"].append(val)

    result["skipped"] = len(result["all"]) - len(result["serials"])
    return result


def check_libs() -> dict:
    status = {
        "python": sys.executable,
        "pillow": False, "pillow_err": "",
        "pyzbar": False, "pyzbar_err": "",
        "zxingcpp": False, "zxingcpp_err": "",
    }
    try:
        from PIL import Image
        status["pillow"] = True
    except ImportError as e:
        status["pillow_err"] = str(e)
    try:
        from pyzbar.pyzbar import decode
        from PIL import Image
        decode(Image.new("RGB", (10, 10)))
        status["pyzbar"] = True
    except ImportError as e:
        status["pyzbar_err"] = f"ImportError: {e}"
    except Exception as e:
        status["pyzbar_err"] = f"Runtime: {e}"
    try:
        import zxingcpp
        status["zxingcpp"] = True
    except ImportError as e:
        status["zxingcpp_err"] = str(e)
    status["any"] = status["pyzbar"] or status["zxingcpp"]
    return status


def libs_status_text(s: dict) -> str:
    lines = [f"🐍 <b>Python:</b> <code>{s['python']}</code>\n"]
    lines.append(("✅" if s["pillow"] else "❌") + " <b>Pillow</b> " + ("" if s["pillow"] else f"<code>{s['pillow_err']}</code>"))
    lines.append(("✅" if s["pyzbar"] else "❌") + " <b>pyzbar</b> " + ("" if s["pyzbar"] else f"<code>{s['pyzbar_err']}</code>"))
    lines.append(("✅" if s["zxingcpp"] else "❌") + " <b>zxing-cpp</b> " + ("" if s["zxingcpp"] else f"<code>{s['zxingcpp_err']}</code>"))
    if not s["any"]:
        lines.append(
            f"\n⚠️ <b>Установите одну из библиотек:</b>\n"
            f"<code>{s['python']} -m pip install zxing-cpp pillow</code>\n"
            f"или: <code>{s['python']} -m pip install pyzbar pillow</code>\n"
            f"На macOS для pyzbar: <code>brew install zbar</code>"
        )
    else:
        lines.append(f"\n✅ Сканер готов — <b>{'zxingcpp' if s['zxingcpp'] else 'pyzbar'}</b>")
    return "\n".join(lines)


# ===========================================================================
#  ОБРАБОТЧИК ПЛОМБ
# ===========================================================================

class PlombProcessor:
    @staticmethod
    def is_valid(plomb: str) -> bool:
        p = plomb.strip()
        return p.isdigit() and len(p) == 10

    @staticmethod
    def split(text) -> list:
        if pd.isna(text):
            return []
        t = str(text).strip()
        if not t:
            return []
        for sep in (",", ";", "\n", "\t", " ", "/", "|", "\\"):
            t = t.replace(sep, ",")
        return [p.strip() for p in t.split(",") if p.strip()]

    def process_file(self, path: str) -> dict:
        try:
            df = pd.read_excel(path, dtype=str, header=None)
            if df.shape[1] < 2:
                raise ValueError("Файл должен содержать минимум 2 столбца")

            ref_plombs: set = set()
            for cell in df.iloc[:, 0]:
                for p in self.split(cell):
                    if self.is_valid(p):
                        ref_plombs.add(p)

            col1_orig, col2_orig, inv, found, not_found = [], [], [], [], []
            total_inv = total_found = total_nf = 0

            for idx in range(len(df)):
                v1 = "" if pd.isna(df.iloc[idx, 0]) else str(df.iloc[idx, 0])
                v2 = "" if pd.isna(df.iloc[idx, 1]) else str(df.iloc[idx, 1])
                col1_orig.append(v1)
                col2_orig.append(v2)
                il, fl, nfl = [], [], []
                for p in self.split(df.iloc[idx, 1]):
                    if not self.is_valid(p):
                        il.append(p); total_inv += 1
                    elif p in ref_plombs:
                        fl.append(p); total_found += 1
                    else:
                        nfl.append(p); total_nf += 1
                inv.append(", ".join(il))
                found.append(", ".join(fl))
                not_found.append(", ".join(nfl))

            out_df = pd.DataFrame({
                "Оригинал_справочник": col1_orig,
                "Оригинал_установленные": col2_orig,
                "Неверный_формат": inv,
                "Найдены_совпадения": found,
                "Не_найдены": not_found,
            })
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            out_path = str(TEMP_DIR / f"результат_обработки_{ts}.xlsx")
            with pd.ExcelWriter(out_path, engine="openpyxl") as w:
                out_df.to_excel(w, index=False, sheet_name="Результат")
                pd.DataFrame({
                    "Показатель": ["Всего строк", "Валидных пломб в справочнике", "Пломб неверного формата", "Пломб с совпадениями", "Пломб без совпадений"],
                    "Значение": [len(df), len(ref_plombs), total_inv, total_found, total_nf],
                }).to_excel(w, index=False, sheet_name="Статистика")
                if total_inv > 0:
                    ex = [{"Строка": i + 1, "Значение": col2_orig[i], "Ошибки": inv[i]} for i in range(min(20, len(inv))) if inv[i]]
                    if ex:
                        pd.DataFrame(ex).to_excel(w, index=False, sheet_name="Примеры_ошибок")
            return {"success": True, "output_path": out_path, "stats": {"total_rows": len(df), "valid_plombs": len(ref_plombs), "invalid": total_inv, "found": total_found, "not_found": total_nf}}
        except Exception as e:
            logger.error(f"PlombProcessor error: {e}")
            return {"success": False, "error": str(e)}


# ===========================================================================
#  ОБРАБОТЧИК ГАРАНТИИ
# ===========================================================================

class GuaranteeProcessor:
    @staticmethod
    def load(path: str) -> pd.DataFrame:
        df = pd.read_excel(path, dtype=str)
        logger.info(f"Файл загружен. Колонки: {df.columns.tolist()}")
        return df

    @staticmethod
    def find_devices(df: pd.DataFrame, devices: list) -> dict:
        result = {"valid": [], "invalid": []}

        status_col_idx = 3
        serial_col_idx = 85
        model_col_idx = 84

        if len(df.columns) <= max(status_col_idx, serial_col_idx, model_col_idx):
            logger.error(f"В файле недостаточно столбцов. Требуется минимум {max(status_col_idx, serial_col_idx, model_col_idx) + 1}")
            return result

        status_col_name = df.columns[status_col_idx]
        serial_col_name = df.columns[serial_col_idx]
        model_col_name = df.columns[model_col_idx]

        logger.info(f"Столбец статуса (D): '{status_col_name}'")
        logger.info(f"Столбец серийных номеров (CH): '{serial_col_name}'")
        logger.info(f"Столбец наименования ПУ (CG): '{model_col_name}'")

        for device_data in devices:
            serial_number = str(device_data["serial"]).strip()
            year = str(device_data["year"]).strip() if device_data["year"] else ""

            logger.info(f"Поиск серийного номера: {serial_number}")

            found_rows = []
            for idx, row in df.iterrows():
                serial_value = str(row[serial_col_name]).strip() if pd.notna(row[serial_col_name]) else ""
                if serial_value == serial_number:
                    found_rows.append((idx, row))

            if not found_rows:
                result["invalid"].append({
                    "serial": serial_number,
                    "year": year,
                    "reason": f"Серийный номер не найден в столбце {serial_col_name}"
                })
                logger.warning(f"Серийный номер {serial_number} не найден")
                continue

            valid_row = None
            invalid_reason = None

            for idx, row in found_rows:
                status_value = str(row[status_col_name]).strip() if pd.notna(row[status_col_name]) else ""
                if "Жалоба" in status_value:
                    valid_row = (idx, row)
                    logger.info(f"Найдена строка с серийным номером {serial_number} и статусом Жалоба")
                    break
                else:
                    invalid_reason = f'Статус "{status_value}" не содержит "Жалоба"'

            if valid_row is not None:
                idx, row = valid_row

                device_info = {
                    "serial_number": serial_number,
                    "year": year,
                    "address_full": "",
                    "request_id": "",
                    "model": "",
                    "status": "Жалоба"
                }

                model_value = str(row[model_col_name]).strip() if pd.notna(row[model_col_name]) else ""
                if not model_value or model_value.lower() == "nan":
                    # Фолбэк: ищем по ключевым словам в названиях столбцов
                    for col in df.columns:
                        col_lower = str(col).lower()
                        if any(kw in col_lower for kw in ["наименование", "модель", "тип прибора", "тип пу", "марка"]):
                            v = str(row[col]).strip() if pd.notna(row[col]) else ""
                            if v and v.lower() != "nan":
                                model_value = v
                                logger.info(f"Модель ПУ найдена по имени столбца '{col}': {v}")
                                break
                if model_value and model_value.lower() != "nan":
                    device_info["model"] = model_value
                    logger.info(f"Модель ПУ: {model_value}")
                else:
                    device_info["model"] = "не указано"
                    logger.warning(f"Модель ПУ не найдена для {serial_number}")

                for col in ["ID заявки", "Номер заявки"]:
                    if col in df.columns and pd.notna(row[col]):
                        device_info["request_id"] = str(row[col]).strip()
                        logger.info(f"ID заявки: {device_info['request_id']}")
                        break

                # Сначала пробуем взять готовый адрес одной строкой (столбец AB)
                address_full = ""
                for col in ["Адрес (одной строкой)", "Адрес"]:
                    if col in df.columns and pd.notna(row[col]):
                        v = str(row[col]).strip()
                        if v and v.lower() != "nan":
                            address_full = v
                            logger.info(f"Адрес (одной строкой): {address_full}")
                            break

                # Если готового адреса нет — собираем из компонентных столбцов
                # AD=Регион, AF=Населенный пункт, AG=Улица, AH=Дом, AI=Корпус, AJ=Квартира
                if not address_full:
                    def _get_col(col_name):
                        if col_name in df.columns and pd.notna(row[col_name]):
                            val = str(row[col_name]).strip()
                            return val if val and val.lower() != "nan" else ""
                        return ""

                    region    = _get_col("Регион")
                    city      = _get_col("Населенный пункт")
                    street    = _get_col("Улица")
                    house     = _get_col("Дом")
                    building  = _get_col("Корпус")
                    apartment = _get_col("Квартира")

                    parts = []
                    if region:
                        parts.append(region)
                    if city:
                        parts.append(city)
                    if street:
                        parts.append(street)
                    if house:
                        parts.append(f"д. {house}")
                    if building:
                        parts.append(f"к. {building}")
                    if apartment:
                        parts.append(f"кв. {apartment}")

                    if parts:
                        address_full = ", ".join(parts)
                        logger.info(f"Адрес собран из компонентов: {address_full}")
                    else:
                        logger.warning(f"Адрес не найден ни в одном столбце для {serial_number}")

                device_info["address_full"] = address_full
                result["valid"].append(device_info)
                logger.info(f"Устройство {serial_number} прошло проверку")

            else:
                result["invalid"].append({
                    "serial": serial_number,
                    "year": year,
                    "reason": invalid_reason or 'Статус не содержит "Жалоба"'
                })
                logger.warning(f"Устройство {serial_number} не прошло проверку по статусу")

        logger.info(f"Итого: валидных — {len(result['valid'])}, невалидных — {len(result['invalid'])}")
        return result

    @staticmethod
    def build_registry(registry_path: str, devices: list) -> str:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out = str(TEMP_DIR / f"реестр_гарантии_{ts}.xlsx")
        df = pd.read_excel(registry_path, dtype=str) if os.path.exists(registry_path) else pd.DataFrame(columns=list("ABCDEFGHIJKLM"))
        for d in devices:
            df = pd.concat([df, pd.DataFrame([{"A": str(len(df) + 1), "B": d.get("model", ""), "C": d.get("serial_number", ""), "D": d.get("year", ""), "E": "", "F": "шт", "G": "1", "H": "", "I": "СЭК", "J": "", "K": "", "L": d.get("address_full", ""), "M": d.get("request_id", "")}])], ignore_index=True)
        df.to_excel(out, index=False)
        return out

    @staticmethod
    def build_acts(template_path: str, devices: list) -> str:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out = str(TEMP_DIR / f"акты_демонтажа_{ts}.docx")
        combined = Document()
        style = combined.styles["Normal"]
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

        def fix_spacing(doc):
            for p in doc.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)

        for i, dev in enumerate(devices):
            doc = Document(template_path)
            fix_spacing(doc)
            repl = {"[[Наименование]]": dev.get("model", "не указано"), "[[Номер]]": dev.get("serial_number", "не указан"), "[[Год]]": dev.get("year", "не указан"), "[[Адрес]]": dev.get("address_full", "не указан")}

            def patch(para):
                for ph, val in repl.items():
                    if ph in para.text:
                        if para.runs:
                            r0 = para.runs[0]
                            fn, fs, bold, ital, ul = r0.font.name, r0.font.size, r0.font.bold, r0.font.italic, r0.font.underline
                            new_text = para.text.replace(ph, val)
                            for run in para.runs: run.text = ""
                            nr = para.add_run(new_text)
                            if fn: nr.font.name = fn
                            if fs: nr.font.size = fs
                            nr.font.bold = bold; nr.font.italic = ital; nr.font.underline = ul
                        else:
                            para.text = para.text.replace(ph, val)

            for p in doc.paragraphs: patch(p)
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs: patch(p)
            try:
                for sec in doc.sections:
                    if sec.header:
                        for p in sec.header.paragraphs: patch(p)
                    if sec.footer:
                        for p in sec.footer.paragraphs: patch(p)
            except Exception as e:
                logger.warning(f"header/footer: {e}")

            for el in doc.element.body.iterchildren():
                combined.element.body.append(el)
            if i < len(devices) - 1:
                combined.add_page_break()

        fix_spacing(combined)
        combined.save(out)
        return out


# ===========================================================================
#  ГЕНЕРАЦИЯ QR-КОДОВ
# ===========================================================================

def generate_qr_pdf(boxes: list) -> str:
    """
    Генерирует PDF с QR-кодами.
    boxes — список словарей: {"label": "Коробка 1", "serials": ["123", "456", ...]}
    Каждый QR на отдельной странице размером 10×10 см — точно под печать этикетки.
    """
    import qrcode
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas as rl_canvas

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = str(TEMP_DIR / f"qr_коды_{ts}.pdf")

    # ── Шрифт с поддержкой кириллицы ──────────────────────────────────────────
    font_name = "Helvetica"
    for font_path in [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("CyrFont", font_path))
                font_name = "CyrFont"
                break
            except Exception:
                pass

    # ── Размер страницы: 10 × 10 cm ───────────────────────────────────────────
    PAGE_W  = 10 * cm
    PAGE_H  = 10 * cm
    MARGIN  = 3 * mm
    INNER_W = PAGE_W - 2 * MARGIN

    # Зоны (снизу вверх — координаты ReportLab)
    SERIAL_H = 0.85 * cm   # блок серийных номеров
    LABEL_H  = 0.72 * cm   # подпись коробки
    GAP      = 1.0 * mm    # зазор между зонами
    QR_SIZE  = PAGE_H - MARGIN - SERIAL_H - GAP - LABEL_H - GAP - MARGIN

    c = rl_canvas.Canvas(out_path, pagesize=(PAGE_W, PAGE_H))

    for i, box in enumerate(boxes):
        label   = box["label"]
        serials = box["serials"]

        # Byte mode — сохраняет переносы строк для числовых данных
        qr_data_bytes = ("\r\n".join(serials)).encode("utf-8")

        qr_obj = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_M,
            box_size=12,
            border=2,
        )
        qr_obj.add_data(qr_data_bytes)
        qr_obj.make(fit=True)
        pil_img = qr_obj.make_image(fill_color="black", back_color="white").convert("RGB")

        tmp_qr = str(TEMP_DIR / f"qr_tmp_{ts}_{i}.png")
        pil_img.save(tmp_qr, "PNG", dpi=(600, 600))

        # 1. Серийные номера — нижняя зона
        c.setFont(font_name, 5.5)
        c.setFillColor(colors.HexColor("#555555"))
        chunks = [serials[j:j+4] for j in range(0, len(serials), 4)]
        lines = ["  |  ".join(chunk) for chunk in chunks]
        line_h = 3.2 * mm
        for li, sline in enumerate(reversed(lines)):
            c.drawCentredString(PAGE_W / 2, MARGIN + li * line_h + 1 * mm, sline)

        # 2. Подпись коробки
        label_y = MARGIN + SERIAL_H + GAP
        c.setFont(font_name, 9)
        c.setFillColor(colors.black)
        c.drawCentredString(PAGE_W / 2, label_y + 1.5 * mm, label)

        # 3. QR-код — верхняя зона, центрирован
        qr_y = MARGIN + SERIAL_H + GAP + LABEL_H + GAP
        qr_x = (PAGE_W - QR_SIZE) / 2
        c.drawImage(tmp_qr, qr_x, qr_y, width=QR_SIZE, height=QR_SIZE,
                    preserveAspectRatio=True, anchor="c")

        c.showPage()

    c.save()

    for i in range(len(boxes)):
        tmp_qr = str(TEMP_DIR / f"qr_tmp_{ts}_{i}.png")
        try:
            os.remove(tmp_qr)
        except Exception:
            pass

    return out_path


def _is_serial_number(value: str) -> bool:
    """
    Строка считается серийным номером, если содержит 5+ цифр подряд.
    """
    import re
    return bool(re.search(r'\d{5,}', value))


def parse_qr_message(text: str) -> dict:
    """
    Парсит сообщение вида:
        Коробка 1
        123456789
        987654321
        ...
    Первая строка считается меткой (label), если она НЕ является серийным номером
    (т.е. не содержит 5+ цифр подряд). Иначе — требуем метку явно.
    Возвращает {"label": "Коробка 1", "serials": ["123456789", ...]} или None при ошибке.
    """
    lines = [l.strip() for l in text.strip().splitlines() if l.strip()]
    if not lines:
        return None

    # Если первая строка НЕ похожа на серийный номер — это метка
    if not _is_serial_number(lines[0]):
        label = lines[0]
        serials = lines[1:]
    else:
        # Все строки — серийные номера, метки нет
        label = ""
        serials = lines

    if not serials:
        return None

    if not label:
        return None

    return {"label": label, "serials": serials}


# ===========================================================================
#  UI HELPERS
# ===========================================================================

plomb_proc = PlombProcessor()
guarantee_proc = GuaranteeProcessor()


def get_template_path(user_id: int) -> Path:
    return USER_TEMPLATES_DIR / f"user_{user_id}_template.docx"


def main_keyboard(user_id: int = 0) -> InlineKeyboardMarkup:
    rows = [
        [
            InlineKeyboardButton("🔍 Обработка пломб", callback_data="plomb"),
            InlineKeyboardButton("📋 Реестр гарантии", callback_data="guarantee"),
        ],
        [
            InlineKeyboardButton("📦 Сканировать штрихкоды", callback_data="scan"),
            InlineKeyboardButton("📎 Загрузить шаблон акта", callback_data="template"),
        ],
        [
            InlineKeyboardButton("📲 Генерация QR-кодов", callback_data="qr"),
            InlineKeyboardButton("🚚 Новая поставка", callback_data="supply"),
        ],
        [InlineKeyboardButton("❓ Помощь", callback_data="help")],
    ]
    if user_id == ADMIN_ID:
        rows.append([InlineKeyboardButton("👑 Панель администратора", callback_data="admin")])
    return InlineKeyboardMarkup(rows)


def back_to_main_btn() -> list:
    return [InlineKeyboardButton("🏠 Главное меню", callback_data="main")]


def scan_after_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📦 Сканировать ещё фото", callback_data="scan")],
        [back_to_main_btn()[0]],
    ])


def cancel_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([[InlineKeyboardButton("❌ Отмена → Главное меню", callback_data="main")]])


def admin_keyboard() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("👥 Список пользователей", callback_data="admin_list")],
        [InlineKeyboardButton("➕ Добавить пользователя", callback_data="admin_add")],
        [back_to_main_btn()[0]],
    ])


def qr_continue_keyboard() -> InlineKeyboardMarkup:
    """Клавиатура после добавления массива QR-кодов."""
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("➕ Продолжить работу", callback_data="qr_continue")],
        [InlineKeyboardButton("✅ Завершить и получить PDF", callback_data="qr_finish")],
        [back_to_main_btn()[0]],
    ])


def _cleanup_session(uid: int):
    sess = user_sessions.pop(uid, {})
    p = sess.get("upload_file", "")
    if p and os.path.exists(p):
        try: os.remove(p)
        except Exception: pass


def _save_prompt(uid: int, msg_id: int):
    """Сохраняет message_id инструкционного сообщения для последующего удаления."""
    if uid not in user_sessions:
        user_sessions[uid] = {}
    user_sessions[uid]["_prompt_msg_id"] = msg_id


async def _delete_prompt(bot, chat_id: int, uid: int):
    """Удаляет сохранённое инструкционное сообщение из чата."""
    msg_id = user_sessions.get(uid, {}).pop("_prompt_msg_id", None)
    if msg_id:
        try:
            await bot.delete_message(chat_id=chat_id, message_id=msg_id)
        except Exception:
            pass


# ===========================================================================
#  ACCESS GUARD
# ===========================================================================

async def check_access(update: Update) -> bool:
    """Проверяет доступ. Если нет — отправляет сообщение и возвращает False."""
    uid = update.effective_user.id
    if is_allowed(uid):
        return True
    text = (
        "⛔ <b>Доступ запрещён</b>\n\n"
        "У вас нет доступа к этому боту.\n"
        f"Ваш Telegram ID: <code>{uid}</code>\n\n"
        "Обратитесь к администратору для получения доступа."
    )
    if update.message:
        await update.message.reply_text(text, parse_mode="HTML")
    elif update.callback_query:
        await update.callback_query.answer("⛔ Нет доступа", show_alert=True)
    return False


# ===========================================================================
#  HANDLERS — ОБЩИЕ
# ===========================================================================

async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await check_access(update):
        return CHOOSING
    uid = update.effective_user.id
    name = update.effective_user.first_name or "пользователь"
    text = (
        f"👋 Привет, <b>{name}</b>!\n\n"
        "🤖 <b>Бот для работы с данными счётчиков</b>\n\n"
        "━━━━━━━━━━━━━━━━━━━━\n"
        "🔍 <b>Пломбы</b> — сравнение двух столбцов Excel\n"
        "📋 <b>Гарантия</b> — реестр и акты демонтажа\n"
        "📦 <b>Штрихкоды</b> — фото → серийные номера\n"
        "📎 <b>Шаблон</b> — загрузить шаблон акта\n"
        "📲 <b>QR-коды</b> — генерация QR по коробкам\n"
        "━━━━━━━━━━━━━━━━━━━━\n\n"
        "Выберите действие:"
    )
    if update.message:
        await update.message.reply_text(text, parse_mode="HTML", reply_markup=main_keyboard(uid))
    elif update.callback_query:
        await update.callback_query.edit_message_text(text, parse_mode="HTML", reply_markup=main_keyboard(uid))
    return CHOOSING


async def cmd_app(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Отправляет кнопку для открытия Mini App."""
    if not await check_access(update):
        return CHOOSING

    keyboard = [
        [KeyboardButton("🚀 Открыть Mini App", web_app=WebAppInfo(url="https://cek-qjva.vercel.app"))],
        [KeyboardButton("🏠 Закрыть (вернуться в бот)")]
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=False)

    text = (
        "📱 <b>Telegram Mini App</b>\n\n"
        "Нажмите кнопку ниже, чтобы открыть веб-интерфейс:\n\n"
        "✨ <b>Доступные функции:</b>\n"
        "• 🔍 Обработка пломб\n"
        "• 📋 Реестр гарантии\n"
        "• 📦 Сканирование штрихкодов\n"
        "• 📎 Загрузка шаблона\n"
        "• 📲 Генерация QR-кодов\n"
        "• 🚚 Новая поставка\n\n"
        "<i>Выберите действие в Mini App — бот автоматически запустит нужный режим.</i>"
    )

    if update.message:
        await update.message.reply_text(text, parse_mode="HTML", reply_markup=reply_markup)
    return CHOOSING


async def web_app_data_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Получает и обрабатывает данные, отправленные из Mini App."""
    if not await check_access(update):
        return CHOOSING

    uid = update.effective_user.id
    data = update.effective_message.web_app_data.data

    # Убираем ReplyKeyboard после получения данных из Mini App
    remove_kb = ReplyKeyboardRemove()

    try:
        webapp_data = json.loads(data)
        action = webapp_data.get("action")

        await update.message.reply_text(
            f"✅ <b>Команда получена из Mini App!</b>\n\n"
            f"Действие: <code>{action}</code>\n\n"
            f"Запускаю соответствующий режим...",
            parse_mode="HTML",
            reply_markup=remove_kb,
        )

        if action == "plomb":
            user_sessions[uid] = {"mode": "plomb"}
            await update.message.reply_text(
                "🔍 <b>Обработка пломб</b>\n\n"
                "Загрузите Excel-файл с двумя столбцами:\n"
                "• Столбец 1 — справочник пломб\n"
                "• Столбец 2 — установленные пломбы",
                parse_mode="HTML",
                reply_markup=cancel_keyboard(),
            )
            return WAITING_FILES

        elif action == "guarantee":
            user_sessions[uid] = {"mode": "guarantee"}
            await update.message.reply_text(
                "📋 <b>Реестр гарантии</b>\n\n"
                "Загрузите файл выгрузки (Excel):",
                parse_mode="HTML",
                reply_markup=cancel_keyboard(),
            )
            return WAITING_FILES

        elif action == "scan":
            user_sessions[uid] = {"mode": "scan"}
            await update.message.reply_text(
                "📦 <b>Сканирование штрихкодов</b>\n\n"
                "Отправьте фото коробки со штрихкодами.\n"
                "<i>Для лучшего качества отправляйте как документ.</i>",
                parse_mode="HTML",
                reply_markup=cancel_keyboard(),
            )
            return WAITING_BARCODE_PHOTO

        elif action == "template":
            user_sessions[uid] = {"mode": "template"}
            await update.message.reply_text(
                "📎 <b>Загрузка шаблона акта</b>\n\n"
                "Отправьте файл <code>.docx</code> с шаблоном.\n"
                "Заполнители: <code>[[Наименование]]</code>, <code>[[Номер]]</code>, "
                "<code>[[Год]]</code>, <code>[[Адрес]]</code>",
                parse_mode="HTML",
                reply_markup=cancel_keyboard(),
            )
            return WAITING_TEMPLATE

        elif action == "qr":
            user_sessions[uid] = {"mode": "qr", "qr_boxes": []}
            await update.message.reply_text(
                "📲 <b>Генерация QR-кодов</b>\n\n"
                "Отправьте сообщение в формате:\n"
                "<code>Коробка 1\n123456789\n987654321</code>\n\n"
                "Первая строка — название коробки, остальные — серийные номера.",
                parse_mode="HTML",
                reply_markup=cancel_keyboard(),
            )
            return WAITING_QR_INPUT

        elif action == "supply":
            return await supply_start(update, context)

        elif action == "help":
            return await cmd_help(update, context)

        else:
            await update.message.reply_text(
                f"❌ Неизвестная команда: <code>{action}</code>",
                parse_mode="HTML",
                reply_markup=main_keyboard(uid),
            )
            return CHOOSING

    except json.JSONDecodeError:
        await update.message.reply_text(
            "❌ Ошибка при обработке данных из Mini App.",
            reply_markup=main_keyboard(uid),
        )
        return CHOOSING
    except Exception as e:
        logger.error(f"web_app_data_handler error: {e}")
        await update.message.reply_text(
            f"❌ Ошибка: {e}",
            reply_markup=main_keyboard(uid),
        )
        return CHOOSING


async def cmd_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await check_access(update):
        return CHOOSING
    text = (
        "📖 <b>Справка по функциям</b>\n\n"
        "🔍 <b>Обработка пломб</b>\n"
        "Загрузите Excel с двумя столбцами пломб.\n"
        "Бот проверит 10-значный формат и найдёт совпадения.\n\n"
        "📦 <b>Сканирование штрихкодов</b>\n"
        "Сфотографируйте коробки со штрихкодами.\n"
        "Бот выдаст только числовые серийные номера (сверху вниз).\n"
        "MAC-адреса и буквенные коды игнорируются.\n"
        "Фото удаляется из чата автоматически.\n\n"
        "📋 <b>Реестр гарантии</b>\n"
        "1. Загрузите файл выгрузки (Excel)\n"
        "2. Введите серийные номера в формате <code>НОМЕР-ГОД</code>\n"
        "3. Получите реестр и акты демонтажа\n\n"
        "📎 <b>Шаблон акта</b>\n"
        "Заполнители: <code>[[Наименование]]</code>, <code>[[Номер]]</code>, <code>[[Год]]</code>, <code>[[Адрес]]</code>\n\n"
        "📲 <b>Генерация QR-кодов</b>\n"
        "Отправьте сообщение в формате:\n"
        "<code>Коробка 1\n123456789\n987654321\n...</code>\n"
        "Первая строка — название коробки, остальные — серийные номера.\n"
        "Нажмите «Продолжить работу» для добавления следующей коробки,\n"
        "«Завершить» — получить PDF со всеми QR-кодами.\n\n"
        "🔧 /checklibs — диагностика библиотек сканирования\n"
        "🔧 /myid — узнать свой Telegram ID"
    )
    kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
    if update.callback_query:
        await update.callback_query.edit_message_text(text, parse_mode="HTML", reply_markup=kb)
    else:
        await update.message.reply_text(text, parse_mode="HTML", reply_markup=kb)


async def cmd_myid(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    await update.message.reply_text(
        f"🆔 Ваш Telegram ID: <code>{uid}</code>\n\nСкопируйте и передайте администратору для получения доступа.",
        parse_mode="HTML"
    )


async def cmd_checklibs(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await check_access(update):
        return
    msg = await update.message.reply_text("🔎 Проверяю библиотеки...")
    s = check_libs()
    await msg.edit_text(
        "🔧 <b>Диагностика библиотек сканирования</b>\n\n" + libs_status_text(s),
        parse_mode="HTML",
    )


async def cmd_cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    _cleanup_session(uid)
    await update.message.reply_text(
        "↩️ Операция отменена.",
        reply_markup=main_keyboard(uid)
    )
    return CHOOSING


# ===========================================================================
#  ADMIN PANEL
# ===========================================================================

async def admin_panel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Показывает панель администратора."""
    q = update.callback_query
    uid = q.from_user.id
    if uid != ADMIN_ID:
        await q.answer("⛔ Только для администратора", show_alert=True)
        return CHOOSING

    await q.answer()
    users = get_all_users()
    count = len(users)
    text = (
        "👑 <b>Панель администратора</b>\n\n"
        f"📊 Активных пользователей: <b>{count}</b>\n\n"
        "Выберите действие:"
    )
    await q.edit_message_text(text, parse_mode="HTML", reply_markup=admin_keyboard())
    return ADMIN_PANEL


async def admin_list_users(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Показывает список пользователей с кнопками удаления."""
    q = update.callback_query
    uid = q.from_user.id
    if uid != ADMIN_ID:
        await q.answer("⛔ Только для администратора", show_alert=True)
        return ADMIN_PANEL

    await q.answer()
    users = get_all_users()

    if not users:
        text = "👥 <b>Список пользователей пуст</b>\n\nДобавьте первого пользователя."
        kb = InlineKeyboardMarkup([
            [InlineKeyboardButton("➕ Добавить пользователя", callback_data="admin_add")],
            [InlineKeyboardButton("◀️ Назад", callback_data="admin")],
        ])
    else:
        text = f"👥 <b>Список пользователей ({len(users)})</b>\n\n"
        rows = []
        for uid_str, info in users.items():
            name = info.get("name", "—")
            added = info.get("added", "—")
            text += f"• <code>{uid_str}</code> — {name}\n  Добавлен: {added}\n\n"
            rows.append([InlineKeyboardButton(
                f"🗑 Удалить {name or uid_str}", callback_data=f"admin_del_{uid_str}"
            )])
        rows.append([InlineKeyboardButton("➕ Добавить пользователя", callback_data="admin_add")])
        rows.append([InlineKeyboardButton("◀️ Назад", callback_data="admin")])
        kb = InlineKeyboardMarkup(rows)

    await q.edit_message_text(text, parse_mode="HTML", reply_markup=kb)
    return ADMIN_PANEL


async def admin_add_prompt(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Запрашивает ID нового пользователя."""
    q = update.callback_query
    uid = q.from_user.id
    if uid != ADMIN_ID:
        await q.answer("⛔ Только для администратора", show_alert=True)
        return ADMIN_PANEL

    await q.answer()
    kb = InlineKeyboardMarkup([[InlineKeyboardButton("◀️ Отмена", callback_data="admin_list")]])
    await q.edit_message_text(
        "➕ <b>Добавить пользователя</b>\n\n"
        "Введите Telegram ID пользователя.\n\n"
        "Пользователь может узнать свой ID командой /myid в этом боте.\n\n"
        "Формат: <code>123456789</code>\n"
        "или с именем: <code>123456789 Иван Иванов</code>",
        parse_mode="HTML",
        reply_markup=kb,
    )
    return ADMIN_ADD_USER


async def admin_receive_user_id(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обрабатывает введённый ID пользователя."""
    uid = update.effective_user.id
    if uid != ADMIN_ID:
        return ADMIN_PANEL

    text = update.message.text.strip()
    parts = text.split(None, 1)

    try:
        new_uid = int(parts[0])
    except ValueError:
        await update.message.reply_text(
            "❌ Неверный формат. Введите числовой ID.\nПример: <code>123456789</code>",
            parse_mode="HTML"
        )
        return ADMIN_ADD_USER

    if new_uid == ADMIN_ID:
        await update.message.reply_text("ℹ️ Это ваш собственный ID — вы уже администратор.")
        return ADMIN_PANEL

    name = parts[1].strip() if len(parts) > 1 else ""
    add_user(new_uid, name)

    await update.message.reply_text(
        f"✅ <b>Пользователь добавлен!</b>\n\n"
        f"🆔 ID: <code>{new_uid}</code>\n"
        f"👤 Имя: {name or '—'}\n\n"
        "Теперь он может использовать бота.",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("👥 К списку пользователей", callback_data="admin_list")],
            [InlineKeyboardButton("➕ Добавить ещё", callback_data="admin_add")],
            [InlineKeyboardButton("🏠 Главное меню", callback_data="main")],
        ])
    )
    return ADMIN_PANEL


async def admin_delete_user(update: Update, context: ContextTypes.DEFAULT_TYPE, del_uid: str):
    """Удаляет пользователя."""
    q = update.callback_query
    uid = q.from_user.id
    if uid != ADMIN_ID:
        await q.answer("⛔ Только для администратора", show_alert=True)
        return ADMIN_PANEL

    await q.answer()
    users = get_all_users()
    user_info = users.get(del_uid, {})
    name = user_info.get("name", del_uid)
    remove_user(int(del_uid))

    await q.answer(f"✅ Пользователь {name} удалён", show_alert=True)

    users = get_all_users()
    if not users:
        text = "👥 <b>Список пользователей пуст</b>"
        kb = InlineKeyboardMarkup([
            [InlineKeyboardButton("➕ Добавить пользователя", callback_data="admin_add")],
            [InlineKeyboardButton("◀️ Назад", callback_data="admin")],
        ])
    else:
        text = f"👥 <b>Список пользователей ({len(users)})</b>\n\n"
        rows = []
        for uid_str, info in users.items():
            n = info.get("name", "—")
            added = info.get("added", "—")
            text += f"• <code>{uid_str}</code> — {n}\n  Добавлен: {added}\n\n"
            rows.append([InlineKeyboardButton(f"🗑 Удалить {n or uid_str}", callback_data=f"admin_del_{uid_str}")])
        rows.append([InlineKeyboardButton("➕ Добавить пользователя", callback_data="admin_add")])
        rows.append([InlineKeyboardButton("◀️ Назад", callback_data="admin")])
        kb = InlineKeyboardMarkup(rows)

    await q.edit_message_text(text, parse_mode="HTML", reply_markup=kb)
    return ADMIN_PANEL


# ===========================================================================
#  HANDLERS — КНОПКИ
# ===========================================================================

async def btn_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()
    uid = q.from_user.id
    data = q.data

    # --- Проверка доступа ---
    if not is_allowed(uid) and data not in ("main",):
        await q.answer("⛔ Нет доступа", show_alert=True)
        return CHOOSING

    # --- Навигация ---
    if data == "main":
        _cleanup_session(uid)
        return await cmd_start(update, context)

    if data == "help":
        return await cmd_help(update, context)

    # --- Админ ---
    if data == "admin":
        return await admin_panel(update, context)

    if data == "admin_list":
        return await admin_list_users(update, context)

    if data == "admin_add":
        return await admin_add_prompt(update, context)

    if data.startswith("admin_del_"):
        del_uid = data.replace("admin_del_", "")
        return await admin_delete_user(update, context, del_uid)

    # --- Режимы работы ---
    if data == "plomb":
        user_sessions[uid] = {"mode": "plomb"}
        kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
        await q.edit_message_text(
            "🔍 <b>Обработка пломб</b>\n\n"
            "Загрузите Excel-файл с двумя столбцами:\n\n"
            "📌 Столбец A — пломбы из справочника\n"
            "📌 Столбец B — установленные пломбы\n\n"
            "Бот сравнит их и выдаст результат.",
            parse_mode="HTML", reply_markup=kb,
        )
        return WAITING_FILES

    if data == "guarantee":
        user_sessions[uid] = {"mode": "guarantee"}
        kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
        await q.edit_message_text(
            "📋 <b>Реестр гарантии — шаг 1 из 3</b>\n\n"
            "Загрузите файл выгрузки (Excel).\n\n"
            "Требования к файлу:\n"
            "• Столбец D — статус (нужно «Жалоба»)\n"
            "• Столбец CH — серийный номер прибора",
            parse_mode="HTML", reply_markup=kb,
        )
        _save_prompt(uid, q.message.message_id)
        return WAITING_FILES

    if data == "scan":
        user_sessions[uid] = {"mode": "scan"}
        kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
        await q.edit_message_text(
            "📦 <b>Сканирование штрихкодов</b>\n\n"
            "Отправьте фото коробки со штрихкодами.\n\n"
            "📌 <b>Советы для лучшего результата:</b>\n"
            "• Хорошее освещение\n"
            "• Держите камеру ровно\n"
            "• Весь лист штрихкодов в кадре\n"
            "• Для максимального качества — отправьте <b>как документ</b> (скрепка → Файл)\n\n"
            "Бот выдаст только числовые серийные номера.\n"
            "Фото автоматически удалится из чата.",
            parse_mode="HTML", reply_markup=kb,
        )
        return WAITING_BARCODE_PHOTO

    if data == "template":
        user_sessions[uid] = {"mode": "template"}
        kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
        await q.edit_message_text(
            "📎 <b>Загрузка шаблона акта демонтажа</b>\n\n"
            "Отправьте файл .docx с вашим шаблоном.\n\n"
            "Шаблон сохранится навсегда для вашего аккаунта.\n\n"
            "📌 <b>Заполнители в шаблоне:</b>\n"
            "• <code>[[Наименование]]</code> — модель прибора\n"
            "• <code>[[Номер]]</code> — серийный номер\n"
            "• <code>[[Год]]</code> — год выпуска\n"
            "• <code>[[Адрес]]</code> — адрес установки",
            parse_mode="HTML", reply_markup=kb,
        )
        _save_prompt(uid, q.message.message_id)
        return WAITING_TEMPLATE

    # --- QR-коды ---
    if data == "qr":
        user_sessions[uid] = {"mode": "qr", "qr_boxes": []}
        kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
        await q.edit_message_text(
            "📲 <b>Генерация QR-кодов</b>\n\n"
            "Отправьте сообщение в следующем формате:\n\n"
            "<code>Коробка 1\n"
            "123456789\n"
            "987654321\n"
            "112233445</code>\n\n"
            "📌 <b>Первая строка</b> — название/номер коробки\n"
            "📌 <b>Остальные строки</b> — серийные номера\n\n"
            "После отправки выберите: продолжить (ещё коробка) или завершить (получить PDF).",
            parse_mode="HTML", reply_markup=kb,
        )
        _save_prompt(uid, q.message.message_id)
        return WAITING_QR_INPUT

    if data == "qr_continue":
        kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
        boxes = user_sessions.get(uid, {}).get("qr_boxes", [])
        await q.edit_message_text(
            f"✅ <b>Добавлено коробок: {len(boxes)}</b>\n\n"
            "Отправьте следующее сообщение с серийными номерами:\n\n"
            "<code>Коробка 2\n"
            "111222333\n"
            "444555666</code>\n\n"
            "📌 Первая строка — название, остальные — серийные номера.",
            parse_mode="HTML", reply_markup=kb,
        )
        _save_prompt(uid, q.message.message_id)
        return WAITING_QR_INPUT

    if data == "qr_finish":
        return await finish_qr(update, context)

    if data in ("process_all", "process_valid", "cancel_op"):
        return await _handle_guarantee_confirm(update, context, data)

    # --- Режим «Новая поставка» ---
    if data == "supply":
        return await supply_start(update, context)

    if data == "supply_photo_ok":
        return await supply_photo_confirmed(update, context)

    if data == "supply_photo_retry":
        return await supply_photo_retry(update, context)

    if data == "supply_edit_serials":
        return await supply_edit_serials_start(update, context)

    if data == "supply_add_more":
        return await supply_add_more(update, context)

    if data == "supply_finish":
        return await supply_finish(update, context)

    return CHOOSING


# ===========================================================================
#  HANDLERS — QR
# ===========================================================================

async def handle_qr_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Принимает текст с серийными номерами для QR-кодов."""
    if not await check_access(update):
        return CHOOSING
    uid = update.effective_user.id
    sess = user_sessions.get(uid, {})

    if sess.get("mode") != "qr":
        # Не в режиме QR — игнорируем
        return CHOOSING

    # Удаляем инструкционное сообщение (qr / qr_continue) и сообщение пользователя
    await _delete_prompt(context.bot, update.message.chat_id, uid)
    try:
        await context.bot.delete_message(
            chat_id=update.message.chat_id, message_id=update.message.message_id
        )
    except Exception:
        pass

    text = update.message.text.strip()
    parsed = parse_qr_message(text)

    if not parsed:
        await update.message.reply_text(
            "❌ <b>Неверный формат.</b>\n\n"
            "Ожидается:\n"
            "<code>Коробка 1\n123456789\n987654321</code>\n\n"
            "Первая строка — название, остальные — серийные номера.",
            parse_mode="HTML",
            reply_markup=cancel_keyboard(),
        )
        return WAITING_QR_INPUT

    if "qr_boxes" not in sess:
        sess["qr_boxes"] = []

    sess["qr_boxes"].append(parsed)
    count = len(sess["qr_boxes"])
    total_serials = sum(len(b["serials"]) for b in sess["qr_boxes"])

    await update.message.reply_text(
        f"✅ <b>Коробка добавлена!</b>\n\n"
        f"📦 Коробок всего: <b>{count}</b>\n"
        f"📋 Серийных номеров всего: <b>{total_serials}</b>\n\n"
        f"<b>Последняя добавленная:</b> {parsed['label']} ({len(parsed['serials'])} номеров)\n\n"
        "Что делаем дальше?",
        parse_mode="HTML",
        reply_markup=qr_continue_keyboard(),
    )
    return WAITING_QR_INPUT


async def finish_qr(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Генерирует PDF с QR-кодами и отправляет пользователю."""
    q = update.callback_query
    uid = q.from_user.id
    sess = user_sessions.get(uid, {})
    boxes = sess.get("qr_boxes", [])

    if not boxes:
        await q.answer("❌ Нет данных для генерации", show_alert=True)
        return WAITING_QR_INPUT

    await q.edit_message_text("⚙️ Генерирую PDF с QR-кодами, подождите...")

    try:
        pdf_path = generate_qr_pdf(boxes)
        ts = datetime.now().strftime("%d.%m.%Y")
        total_serials = sum(len(b["serials"]) for b in boxes)

        with open(pdf_path, "rb") as fh:
            await context.bot.send_document(
                chat_id=q.message.chat_id,
                document=fh,
                filename=f"QR_коды_{ts}.pdf",
                caption=(
                    f"📲 <b>QR-коды готовы!</b>\n\n"
                    f"📦 Коробок: {len(boxes)}\n"
                    f"📋 Серийных номеров: {total_serials}\n\n"
                    "Каждый QR-код на отдельной странице."
                ),
                parse_mode="HTML",
            )

        try:
            os.remove(pdf_path)
        except Exception:
            pass

        _cleanup_session(uid)
        await context.bot.send_message(
            chat_id=q.message.chat_id,
            text="Выберите следующее действие:",
            reply_markup=main_keyboard(uid),
        )
        return CHOOSING

    except ImportError as e:
        await q.edit_message_text(
            f"❌ <b>Не установлены библиотеки для генерации QR:</b>\n\n"
            f"<code>{e}</code>\n\n"
            "Установите:\n"
            "<code>pip install qrcode reportlab pillow</code>",
            parse_mode="HTML",
        )
        await context.bot.send_message(
            chat_id=q.message.chat_id,
            text="Выберите действие:",
            reply_markup=main_keyboard(uid),
        )
        return CHOOSING
    except Exception as e:
        logger.error(f"finish_qr error: {e}")
        import traceback; logger.error(traceback.format_exc())
        await q.edit_message_text(f"❌ Ошибка при генерации PDF: {e}")
        await context.bot.send_message(
            chat_id=q.message.chat_id,
            text="Выберите действие:",
            reply_markup=main_keyboard(uid),
        )
        return CHOOSING


# ===========================================================================
#  HANDLERS — ФАЙЛЫ
# ===========================================================================

async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await check_access(update):
        return CHOOSING
    uid = update.effective_user.id
    mode = user_sessions.get(uid, {}).get("mode")
    if not mode:
        await update.message.reply_text("Выберите режим через /start", reply_markup=main_keyboard(uid))
        return CHOOSING

    doc = update.message.document
    if not doc.file_name.lower().endswith((".xlsx", ".xls")):
        kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
        await update.message.reply_text("❌ Нужен файл Excel (.xlsx или .xls)", reply_markup=kb)
        return WAITING_FILES

    status = await update.message.reply_text("📥 Получаю файл...")
    # Удаляем сообщение пользователя с файлом
    try:
        await context.bot.delete_message(
            chat_id=update.message.chat_id, message_id=update.message.message_id
        )
    except Exception:
        pass
    try:
        f = await context.bot.get_file(doc.file_id)
        path = str(TEMP_DIR / f"tmp_{uid}_{doc.file_name}")
        await f.download_to_drive(path)

        if mode == "plomb":
            await status.edit_text("⚙️ Обрабатываю пломбы...")
            res = plomb_proc.process_file(path)
            if res["success"]:
                st = res["stats"]
                await status.edit_text(
                    "✅ <b>Обработка завершена!</b>\n\n"
                    f"📊 <b>Статистика:</b>\n"
                    f"• Всего строк: {st['total_rows']}\n"
                    f"• Пломб в справочнике: {st['valid_plombs']}\n"
                    f"• Неверный формат: {st['invalid']}\n"
                    f"• Совпадений: {st['found']}\n"
                    f"• Без совпадений: {st['not_found']}",
                    parse_mode="HTML",
                )
                with open(res["output_path"], "rb") as fh:
                    await update.message.reply_document(
                        document=fh,
                        filename=f"Пломбы_{datetime.now().strftime('%d.%m.%Y')}.xlsx",
                        caption="📎 Результат обработки пломб",
                    )
                for p in (path, res["output_path"]):
                    try: os.remove(p)
                    except Exception: pass
            else:
                await status.edit_text(f"❌ Ошибка: <code>{res['error']}</code>", parse_mode="HTML")
                try: os.remove(path)
                except Exception: pass
            await update.message.reply_text("Выберите следующее действие:", reply_markup=main_keyboard(uid))
            return CHOOSING

        elif mode == "guarantee":
            user_sessions[uid]["upload_file"] = path
            await status.edit_text("✅ Файл загружен.")
            # Удаляем сообщение шага 1 (инструкция с требованиями к файлу)
            await _delete_prompt(context.bot, update.message.chat_id, uid)
            kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])
            msg2 = await update.message.reply_text(
                "📋 <b>Шаг 2 из 3 — Серийные номера</b>\n\n"
                "Введите серийные номера приборов.\n\n"
                "📌 Формат каждой строки:\n"
                "<code>СЕРИЙНЫЙ_НОМЕР-ГОД</code>\n\n"
                "Пример:\n"
                "<code>021241589506-2024\n9876543210-2023</code>\n\n"
                "Каждый номер — с новой строки.",
                parse_mode="HTML", reply_markup=kb,
            )
            # Сохраняем ID сообщения "Файл загружен" и шага 2 для удаления
            user_sessions[uid]["_step_msgs"] = [status.message_id, msg2.message_id]
            return WAITING_SERIALS

    except Exception as e:
        logger.error(f"handle_file error: {e}")
        await status.edit_text(f"❌ Ошибка: {e}")
        await update.message.reply_text("Выберите действие:", reply_markup=main_keyboard(uid))
        return CHOOSING


async def handle_serials(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await check_access(update):
        return CHOOSING
    uid = update.effective_user.id
    sess = user_sessions.get(uid, {})
    if "upload_file" not in sess:
        await update.message.reply_text("Сначала загрузите файл через /start")
        return WAITING_FILES

    status = await update.message.reply_text("🔄 Проверяю серийные номера...")
    # Удаляем сообщение пользователя с серийными номерами
    try:
        await context.bot.delete_message(
            chat_id=update.message.chat_id, message_id=update.message.message_id
        )
    except Exception:
        pass
    # Удаляем сообщения шага 2 ("✅ Файл загружен." и инструкцию с серийными номерами)
    for mid in user_sessions.get(uid, {}).pop("_step_msgs", []):
        try:
            await context.bot.delete_message(chat_id=update.message.chat_id, message_id=mid)
        except Exception:
            pass
    try:
        devices = []
        for line in update.message.text.strip().splitlines():
            line = line.strip()
            if not line:
                continue
            if "-" in line:
                sn, yr = line.split("-", 1)
            else:
                sn, yr = line, ""
            devices.append({"serial": sn.strip(), "year": yr.strip()})

        if not devices:
            await status.edit_text("❌ Не найдено серийных номеров.")
            return WAITING_SERIALS

        df = guarantee_proc.load(sess["upload_file"])
        val = guarantee_proc.find_devices(df, devices)
        sess["validation"] = val
        sess["all_devices"] = devices

        lines = ["🔍 <b>Результаты проверки — шаг 3 из 3</b>\n"]
        if val["valid"]:
            lines.append(f"✅ <b>Прошли проверку ({len(val['valid'])}):</b>")
            for d in val["valid"]:
                lines.append(
                    f"  • <code>{d['serial_number']}</code> ({d['year']})\n"
                    f"    Модель: {d['model']}\n"
                    f"    Адрес: {d['address_full'] or 'не найден'}"
                )
        if val["invalid"]:
            lines.append(f"\n❌ <b>Не прошли проверку ({len(val['invalid'])}):</b>")
            for d in val["invalid"]:
                lines.append(f"  • <code>{d['serial']}</code> ({d['year']}) — {d['reason']}")

        lines.append(f"\n📊 Итого: {len(val['valid'])} ✅  {len(val['invalid'])} ❌\n\nВыберите вариант обработки:")

        kb = InlineKeyboardMarkup([
            [
                InlineKeyboardButton(f"✅ Все ({len(devices)})", callback_data="process_all"),
                InlineKeyboardButton(f"✓ Только верные ({len(val['valid'])})", callback_data="process_valid"),
            ],
            [InlineKeyboardButton("◀️ Отмена → Главное меню", callback_data="cancel_op")],
        ])
        await status.edit_text("\n".join(lines), parse_mode="HTML", reply_markup=kb)
        return CONFIRM_DATA

    except Exception as e:
        logger.error(f"handle_serials error: {e}")
        await status.edit_text(f"❌ Ошибка: {e}")
        return WAITING_SERIALS


async def _handle_guarantee_confirm(update: Update, context: ContextTypes.DEFAULT_TYPE, action: str):
    q = update.callback_query
    uid = q.from_user.id
    sess = user_sessions.get(uid, {})

    if action == "cancel_op":
        _cleanup_session(uid)
        return await cmd_start(update, context)

    await q.edit_message_text("⚙️ Создаю документы, подождите...")
    try:
        val = sess.get("validation", {"valid": [], "invalid": []})
        if action == "process_all":
            to_proc = val["valid"] + [
                {"serial_number": d["serial"], "year": d["year"], "model": "НЕ ПРОШЕЛ ПРОВЕРКУ", "address_full": "", "request_id": ""}
                for d in val["invalid"]
            ]
            label = f"все {len(to_proc)} приборов"
        else:
            to_proc = val["valid"]
            label = f"{len(to_proc)} валидных приборов"

        if not to_proc:
            await q.edit_message_text("❌ Нет приборов для обработки.")
            await context.bot.send_message(chat_id=q.message.chat_id, text="Выберите действие:", reply_markup=main_keyboard(uid))
            return CHOOSING

        tpl = get_template_path(uid)
        if not tpl.exists():
            await q.edit_message_text(
                "❌ <b>Шаблон акта не найден!</b>\n\nЗагрузите шаблон через главное меню.",
                parse_mode="HTML"
            )
            await context.bot.send_message(chat_id=q.message.chat_id, text="Выберите действие:", reply_markup=main_keyboard(uid))
            return CHOOSING

        reg_tpl = TEMPLATES_DIR / "реестр_гарантии_шаблон.xlsx"
        if not reg_tpl.exists():
            pd.DataFrame(columns=list("ABCDEFGHIJKLM")).to_excel(reg_tpl, index=False)

        reg_path = guarantee_proc.build_registry(str(reg_tpl), to_proc)
        acts_path = guarantee_proc.build_acts(str(tpl), to_proc)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_path = str(TEMP_DIR / f"документы_{ts}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(reg_path, f"реестр_{ts}.xlsx")
            zf.write(acts_path, f"акты_{len(to_proc)}шт_{ts}.docx")

        with open(zip_path, "rb") as fh:
            await context.bot.send_document(
                chat_id=q.message.chat_id, document=fh,
                filename=f"Документы_{ts}.zip",
                caption=(
                    f"📦 <b>Документы готовы!</b>\n\n"
                    f"Обработано: {label}\n"
                    f"✅ Валидных: {len(val['valid'])}\n"
                    f"❌ Невалидных: {len(val['invalid'])}\n\n"
                    f"📄 В архиве:\n"
                    f"• Реестр гарантии (.xlsx)\n"
                    f"• Акты демонтажа (.docx)"
                ),
                parse_mode="HTML",
            )

        await q.edit_message_text("✅ <b>Документы созданы и отправлены!</b>", parse_mode="HTML")

        for p in (reg_path, acts_path, zip_path):
            try: os.remove(p)
            except Exception: pass
        _cleanup_session(uid)

    except Exception as e:
        logger.error(f"guarantee_confirm error: {e}")
        import traceback; logger.error(traceback.format_exc())
        await q.edit_message_text(f"❌ Ошибка при создании документов: {e}")

    await context.bot.send_message(
        chat_id=q.message.chat_id,
        text="Выберите следующее действие:",
        reply_markup=main_keyboard(uid)
    )
    return CHOOSING


# ===========================================================================
#  HANDLER — ШАБЛОН
# ===========================================================================

async def handle_template_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await check_access(update):
        return CHOOSING
    uid = update.effective_user.id
    doc = update.message.document
    if not doc.file_name.lower().endswith(".docx"):
        await update.message.reply_text("❌ Нужен файл .docx", reply_markup=cancel_keyboard())
        return WAITING_TEMPLATE

    status = await update.message.reply_text("📥 Сохраняю шаблон...")
    # Удаляем инструкцию «Загрузка шаблона» и файл пользователя
    await _delete_prompt(context.bot, update.message.chat_id, uid)
    try:
        await context.bot.delete_message(
            chat_id=update.message.chat_id, message_id=update.message.message_id
        )
    except Exception:
        pass
    try:
        f = await context.bot.get_file(doc.file_id)
        tpl_path = get_template_path(uid)
        await f.download_to_drive(str(tpl_path))

        if not DEFAULT_TEMPLATE_PATH.exists():
            shutil.copy2(str(tpl_path), str(DEFAULT_TEMPLATE_PATH))

        d = Document(str(tpl_path))
        full_text = " ".join(p.text for p in d.paragraphs)
        missing = [ph for ph in ["[[Наименование]]", "[[Номер]]", "[[Год]]", "[[Адрес]]"] if ph not in full_text]
        if missing:
            await update.message.reply_text(
                "⚠️ <b>Внимание!</b> В шаблоне не найдены заполнители:\n" + "\n".join(f"• {m}" for m in missing),
                parse_mode="HTML",
            )

        await status.edit_text(
            f"✅ <b>Шаблон сохранён!</b>\n\nФайл: <code>{doc.file_name}</code>",
            parse_mode="HTML"
        )
        user_sessions.pop(uid, None)
        await update.message.reply_text("Выберите действие:", reply_markup=main_keyboard(uid))
        return CHOOSING
    except Exception as e:
        await status.edit_text(f"❌ Ошибка: {e}")
        return CHOOSING


# ===========================================================================
#  МОДУЛЬ «НОВАЯ ПОСТАВКА» — сканирование + генерация QR за один проход
# ===========================================================================
#
# Сессионные данные хранятся в user_sessions[uid]["supply"]:
#   current_photo_codes — список штрихкодов с последнего обработанного фото
#   current_box_number  — номер коробки, введённый пользователем
#   processed_boxes     — счётчик обработанных коробок за сессию
#   qr_boxes            — список коробок для generate_qr_pdf
#                         [{"label": "...", "serials": [...]}, ...]
# ===========================================================================


def _supply_session(uid: int) -> dict:
    """Возвращает (и при необходимости создаёт) словарь сессии поставки."""
    if uid not in user_sessions:
        user_sessions[uid] = {}
    if "supply" not in user_sessions[uid]:
        user_sessions[uid]["supply"] = {
            "current_photo_codes": [],
            "current_box_number": "",
            "processed_boxes": 0,
            "qr_boxes": [],
        }
    return user_sessions[uid]["supply"]


def _supply_clear_current(uid: int):
    """Сбрасывает данные ТЕКУЩЕЙ коробки, не трогая счётчик и накопленный список."""
    s = _supply_session(uid)
    s["current_photo_codes"] = []
    s["current_box_number"] = ""


async def supply_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Точка входа: кнопка «🚚 Новая поставка» или /supply."""
    if not await check_access(update):
        return CHOOSING

    uid = update.effective_user.id
    # Инициализируем чистую сессию поставки
    user_sessions[uid] = {"mode": "supply", "supply": {
        "current_photo_codes": [],
        "current_box_number": "",
        "processed_boxes": 0,
        "qr_boxes": [],
    }}

    text = (
        "🚚 <b>Режим «Новая поставка»</b>\n\n"
        "Я помогу отсканировать штрихкоды с коробок и сгенерировать QR-коды.\n\n"
        "📸 <b>Отправьте фото коробки</b>, чтобы были видны все штрихкоды.\n\n"
        "💡 <i>Совет: отправляйте фото <b>как документ</b> (скрепка → Файл) "
        "— качество распознавания выше.</i>\n\n"
        "/cancel — выйти из режима поставки"
    )
    kb = InlineKeyboardMarkup([[back_to_main_btn()[0]]])

    q = update.callback_query
    if q:
        await q.edit_message_text(text, parse_mode="HTML", reply_markup=kb)
        _save_prompt(uid, q.message.message_id)
    else:
        sent = await update.message.reply_text(text, parse_mode="HTML", reply_markup=kb)
        _save_prompt(uid, sent.message_id)

    return SUPPLY_WAITING_PHOTO


async def handle_supply_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Шаг 1: получаем фото, сканируем штрихкоды, показываем результат."""
    if not await check_access(update):
        return SUPPLY_WAITING_PHOTO

    uid = update.effective_user.id
    msg = update.message

    photo = msg.photo[-1] if msg.photo else None
    doc = (
        msg.document
        if msg.document and msg.document.mime_type
           and msg.document.mime_type.startswith("image/")
        else None
    )

    if not photo and not doc:
        await msg.reply_text(
            "❌ Ожидаю фото или изображение как документ. Попробуйте ещё раз.",
            reply_markup=cancel_keyboard(),
        )
        return SUPPLY_WAITING_PHOTO

    # Удаляем инструкционное сообщение «Отправьте фото»
    await _delete_prompt(context.bot, msg.chat_id, uid)

    status = await msg.reply_text("🔍 Сканирую штрихкоды на фото...")

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    img_path = str(TEMP_DIR / f"supply_{uid}_{ts}.jpg")

    try:
        tg_file = await context.bot.get_file(doc.file_id if doc else photo.file_id)
        await tg_file.download_to_drive(img_path)

        # Используем существующую функцию сканирования
        result = scan_barcodes(img_path)

        if not result["ok"]:
            # Библиотека не найдена
            s = check_libs()
            await status.edit_text(
                "❌ <b>Библиотека сканирования не найдена!</b>\n\n"
                + libs_status_text(s),
                parse_mode="HTML",
            )
            await msg.reply_text(
                "Отправьте фото ещё раз или /cancel для выхода.",
                reply_markup=cancel_keyboard(),
            )
            return SUPPLY_WAITING_PHOTO

        # Берём ТОЛЬКО числовые серийные номера (MAC-адреса и нечисловые коды игнорируются)
        serial_codes = result.get("serials", [])
        all_codes = result.get("all", [])
        skipped = result.get("skipped", 0)

        # Если нет вообще ничего — просим переснять
        if not all_codes:
            await status.edit_text(
                "⚠️ <b>Штрихкоды не обнаружены.</b>\n\n"
                "Советы:\n"
                "• Хорошее освещение, без теней\n"
                "• Штрихкоды чёткие, не смазанные\n"
                "• Отправьте фото <b>как документ</b> (без сжатия)\n"
                "• Подойдите ближе к коробке",
                parse_mode="HTML",
            )
            await msg.reply_text(
                "Отправьте фото ещё раз или /cancel для выхода.",
                reply_markup=cancel_keyboard(),
            )
            return SUPPLY_WAITING_PHOTO

        # Если все штрихкоды нечисловые (только MAC-адреса) — сообщаем об этом
        if not serial_codes:
            await status.edit_text(
                f"⚠️ <b>Считано штрихкодов: {len(all_codes)}</b>\n\n"
                "Все они нечисловые (MAC-адреса и т.п.).\n"
                "<b>Серийных номеров (числовых) не найдено.</b>\n\n"
                "Попробуйте отправить фото <b>как документ</b> для лучшего качества "
                "или убедитесь, что числовой штрихкод попал в кадр.",
                parse_mode="HTML",
            )
            await msg.reply_text(
                "Отправьте фото ещё раз или /cancel для выхода.",
                reply_markup=cancel_keyboard(),
            )
            return SUPPLY_WAITING_PHOTO

        # Числовые серийные номера найдены — сохраняем в сессии
        _supply_session(uid)["current_photo_codes"] = serial_codes

        # Формируем нумерованный список для показа пользователю
        codes_list = "\n".join(f"{i+1}. <code>{c}</code>" for i, c in enumerate(serial_codes))
        skipped_note = f"\n\n<i>Нечисловых пропущено (MAC и др.): {skipped}</i>" if skipped else ""
        await status.edit_text(
            f"✅ <b>Серийных номеров найдено: {len(serial_codes)}</b>{skipped_note}\n\n"
            f"{codes_list}\n\n"
            "Всё верно?",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([
                [
                    InlineKeyboardButton("✅ Да, всё верно", callback_data="supply_photo_ok"),
                    InlineKeyboardButton("❌ Нет, переснять", callback_data="supply_photo_retry"),
                ],
                [
                    InlineKeyboardButton("✏️ Редактировать список", callback_data="supply_edit_serials"),
                ],
            ]),
        )

        # Удаляем исходное фото из чата (не обязательно, но по аналогии с существующим скан-режимом)
        try:
            await context.bot.delete_message(chat_id=msg.chat_id, message_id=msg.message_id)
        except Exception:
            pass

    except Exception as e:
        logger.error(f"handle_supply_photo error: {e}")
        import traceback; logger.error(traceback.format_exc())
        await status.edit_text(f"❌ Ошибка при сканировании: {e}")
        await msg.reply_text("Попробуйте ещё раз или /cancel.", reply_markup=cancel_keyboard())

    finally:
        if os.path.exists(img_path):
            try:
                os.remove(img_path)
            except Exception:
                pass

    # Остаёмся в состоянии SUPPLY_WAITING_PHOTO — ждём нажатия inline-кнопки
    return SUPPLY_WAITING_PHOTO


async def supply_edit_serials_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Пользователь нажал «Редактировать список» — просим ввести серийники вручную."""
    if not await check_access(update):
        return SUPPLY_WAITING_PHOTO

    uid = update.effective_user.id
    q = update.callback_query
    await q.answer()

    supply = _supply_session(uid)
    current = supply.get("current_photo_codes", [])
    current_text = "\n".join(current)

    await q.edit_message_text(
        "✏️ <b>Редактирование серийных номеров</b>\n\n"
        "Отправьте исправленный список — <b>каждый номер с новой строки</b>.\n"
        "Можно добавить пропущенные, удалить лишние или исправить ошибочные.\n\n"
        "<b>Текущий список:</b>\n"
        f"<code>{current_text}</code>\n\n"
        "✍️ Отправьте исправленный список:",
        parse_mode="HTML",
        reply_markup=cancel_keyboard(),
    )
    _save_prompt(uid, q.message.message_id)
    return SUPPLY_EDITING_SERIALS


async def handle_supply_edit_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Получаем отредактированный список серийников от пользователя."""
    if not await check_access(update):
        return SUPPLY_EDITING_SERIALS

    uid = update.effective_user.id
    msg = update.message

    # Удаляем инструкционное сообщение и сообщение пользователя
    await _delete_prompt(context.bot, msg.chat_id, uid)
    try:
        await context.bot.delete_message(chat_id=msg.chat_id, message_id=msg.message_id)
    except Exception:
        pass

    raw_lines = [l.strip() for l in msg.text.strip().splitlines() if l.strip()]
    serials = [l for l in raw_lines if _is_numeric_serial(l)]
    non_serials = [l for l in raw_lines if not _is_numeric_serial(l)]

    if not serials:
        await msg.reply_text(
            "❌ <b>Не найдено ни одного числового серийного номера.</b>\n\n"
            "Каждая строка должна содержать только цифры (≥6 знаков).\n"
            "Попробуйте ещё раз или нажмите /cancel.",
            parse_mode="HTML",
            reply_markup=cancel_keyboard(),
        )
        return SUPPLY_EDITING_SERIALS

    _supply_session(uid)["current_photo_codes"] = serials

    codes_list = "\n".join(f"{i+1}. <code>{c}</code>" for i, c in enumerate(serials))
    warn = ""
    if non_serials:
        warn = f"\n\n⚠️ <i>Пропущены нечисловые строки: {', '.join(non_serials[:5])}</i>"

    await msg.reply_text(
        f"✅ <b>Список обновлён: {len(serials)} серийных номеров</b>{warn}\n\n"
        f"{codes_list}\n\n"
        "Всё верно?",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("✅ Да, всё верно", callback_data="supply_photo_ok"),
                InlineKeyboardButton("❌ Нет, переснять", callback_data="supply_photo_retry"),
            ],
            [
                InlineKeyboardButton("✏️ Редактировать ещё раз", callback_data="supply_edit_serials"),
            ],
        ]),
    )
    return SUPPLY_WAITING_PHOTO


async def supply_edit_serials_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Пользователь нажал «Редактировать список» — просим ввести серийники вручную."""
    if not await check_access(update):
        return SUPPLY_WAITING_PHOTO

    uid = update.effective_user.id
    q = update.callback_query
    await q.answer()

    supply = _supply_session(uid)
    current = supply.get("current_photo_codes", [])
    current_text = "\n".join(current)

    await q.edit_message_text(
        "✏️ <b>Редактирование серийных номеров</b>\n\n"
        "Отправьте исправленный список — <b>каждый номер с новой строки</b>.\n"
        "Можно добавить пропущенные, удалить лишние или исправить ошибочные.\n\n"
        "<b>Текущий список:</b>\n"
        f"<code>{current_text}</code>\n\n"
        "✍️ Отправьте исправленный список:",
        parse_mode="HTML",
        reply_markup=cancel_keyboard(),
    )
    _save_prompt(uid, q.message.message_id)
    return SUPPLY_EDITING_SERIALS


async def handle_supply_edit_input(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Получаем отредактированный список серийников от пользователя."""
    if not await check_access(update):
        return SUPPLY_EDITING_SERIALS

    uid = update.effective_user.id
    msg = update.message

    # Удаляем инструкционное сообщение и сообщение пользователя
    await _delete_prompt(context.bot, msg.chat_id, uid)
    try:
        await context.bot.delete_message(chat_id=msg.chat_id, message_id=msg.message_id)
    except Exception:
        pass

    raw_lines = [l.strip() for l in msg.text.strip().splitlines() if l.strip()]
    serials = [l for l in raw_lines if _is_numeric_serial(l)]
    non_serials = [l for l in raw_lines if not _is_numeric_serial(l)]

    if not serials:
        await msg.reply_text(
            "❌ <b>Не найдено ни одного числового серийного номера.</b>\n\n"
            "Каждая строка должна содержать только цифры (≥6 знаков).\n"
            "Попробуйте ещё раз или нажмите /cancel.",
            parse_mode="HTML",
            reply_markup=cancel_keyboard(),
        )
        return SUPPLY_EDITING_SERIALS

    # Обновляем серийные номера в сессии
    _supply_session(uid)["current_photo_codes"] = serials

    codes_list = "\n".join(f"{i+1}. <code>{c}</code>" for i, c in enumerate(serials))
    warn = ""
    if non_serials:
        warn = f"\n\n⚠️ <i>Пропущены нечисловые строки: {', '.join(non_serials[:5])}</i>"

    await msg.reply_text(
        f"✅ <b>Список обновлён: {len(serials)} серийных номеров</b>{warn}\n\n"
        f"{codes_list}\n\n"
        "Всё верно?",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("✅ Да, всё верно", callback_data="supply_photo_ok"),
                InlineKeyboardButton("❌ Нет, переснять", callback_data="supply_photo_retry"),
            ],
            [
                InlineKeyboardButton("✏️ Редактировать ещё раз", callback_data="supply_edit_serials"),
            ],
        ]),
    )
    return SUPPLY_WAITING_PHOTO


async def supply_photo_confirmed(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Шаг 2 → 3: пользователь подтвердил фото, просим ввести номер коробки."""
    if not await check_access(update):
        return SUPPLY_WAITING_PHOTO

    uid = update.effective_user.id
    q = update.callback_query
    await q.answer()

    supply = _supply_session(uid)
    codes = supply.get("current_photo_codes", [])

    if not codes:
        await q.edit_message_text(
            "⚠️ Нет данных о серийных номерах. Отправьте фото ещё раз.",
            reply_markup=cancel_keyboard(),
        )
        return SUPPLY_WAITING_PHOTO

    await q.edit_message_text(
        f"✅ Отлично! Зафиксировано <b>{len(codes)}</b> серийных номеров.\n\n"
        "✏️ <b>Введите номер этой коробки</b>\n"
        "(например: <code>BOX-101</code> или <code>Палета №5</code>):",
        parse_mode="HTML",
        reply_markup=cancel_keyboard(),
    )
    _save_prompt(uid, q.message.message_id)
    return SUPPLY_WAITING_NUMBER


async def supply_photo_retry(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Шаг 2 → назад: пользователь хочет переснять фото."""
    uid = update.effective_user.id
    q = update.callback_query
    await q.answer()

    _supply_clear_current(uid)

    await q.edit_message_text(
        "📸 <b>Отправьте новое фото коробки</b>, чтобы были видны все штрихкоды.\n\n"
        "/cancel — выйти из режима поставки",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup([[back_to_main_btn()[0]]]),
    )
    # Сохраняем id — handle_supply_photo удалит это сообщение когда придёт новое фото
    _save_prompt(uid, q.message.message_id)
    return SUPPLY_WAITING_PHOTO


async def handle_supply_number(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Шаг 3: получаем номер коробки, сохраняем в сессии и предлагаем добавить ещё или завершить."""
    if not await check_access(update):
        return SUPPLY_WAITING_NUMBER

    uid = update.effective_user.id
    msg = update.message
    box_number = msg.text.strip()

    if not box_number:
        await msg.reply_text(
            "❌ Номер коробки не может быть пустым. Введите ещё раз:",
            reply_markup=cancel_keyboard(),
        )
        return SUPPLY_WAITING_NUMBER

    supply = _supply_session(uid)
    codes = supply.get("current_photo_codes", [])

    if not codes:
        await msg.reply_text(
            "⚠️ Не найдены серийные номера для этой коробки. Начните заново — отправьте фото.",
            reply_markup=cancel_keyboard(),
        )
        return SUPPLY_WAITING_PHOTO

    # Удаляем сообщение «Введите номер коробки» и сообщение пользователя
    await _delete_prompt(context.bot, msg.chat_id, uid)
    try:
        await context.bot.delete_message(chat_id=msg.chat_id, message_id=msg.message_id)
    except Exception:
        pass

    supply["current_box_number"] = box_number

    # Одна запись на всю коробку: один QR-код содержит ВСЕ серийные номера коробки
    new_qr_entry = {"label": box_number, "serials": codes}

    # Добавляем в накопленный список сессии (PDF не генерируем — отправим всё в конце)
    supply["qr_boxes"].append(new_qr_entry)
    supply["processed_boxes"] += 1
    processed = supply["processed_boxes"]
    total_qr = len(supply["qr_boxes"])  # = количество коробок = количество QR-кодов

    # Список серийных номеров текущей коробки для показа
    codes_preview = "\n".join(f"  • <code>{c}</code>" for c in codes)

    await msg.reply_text(
        f"✅ <b>Коробка «{box_number}» добавлена!</b>\n\n"
        f"<b>Серийные номера ({len(codes)} шт.):</b>\n{codes_preview}\n"
        f"<i>→ будут закодированы в один QR-код</i>\n\n"
        f"📊 Всего коробок за сессию: <b>{processed}</b>  (= {total_qr} QR-кода в PDF)\n\n"
        "Что делаем дальше?",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup([
            [InlineKeyboardButton("➕ Добавить ещё коробку", callback_data="supply_add_more")],
            [InlineKeyboardButton("🏁 Завершить приёмку и получить PDF", callback_data="supply_finish")],
        ]),
    )

    # Удаляем сообщение пользователя с номером коробки
    try:
        await context.bot.delete_message(chat_id=msg.chat_id, message_id=msg.message_id)
    except Exception:
        pass

    # Очищаем данные текущей коробки (штрихкоды + номер)
    _supply_clear_current(uid)

    # Остаёмся в состоянии ожидания следующего действия (inline-кнопки)
    return SUPPLY_WAITING_NUMBER


async def supply_add_more(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Пользователь хочет добавить ещё одну коробку — возвращаемся к фото."""
    uid = update.effective_user.id
    q = update.callback_query
    await q.answer()

    supply = _supply_session(uid)
    processed = supply.get("processed_boxes", 0)
    total_qr = len(supply.get("qr_boxes", []))

    await q.edit_message_text(
        f"📸 <b>Отправьте фото следующей коробки</b>\n\n"
        f"Уже обработано коробок: <b>{processed}</b>  |  QR-кодов в PDF будет: <b>{total_qr}</b>\n\n"
        "💡 <i>Отправляйте фото <b>как документ</b> для лучшего распознавания.</i>\n\n"
        "/cancel — выйти из режима поставки",
        parse_mode="HTML",
        reply_markup=InlineKeyboardMarkup([[back_to_main_btn()[0]]]),
    )
    _save_prompt(uid, q.message.message_id)
    return SUPPLY_WAITING_PHOTO


async def supply_finish(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Пользователь завершил приёмку — генерируем итоговый PDF со всеми QR-кодами и отправляем."""
    uid = update.effective_user.id
    q = update.callback_query
    await q.answer()

    supply = _supply_session(uid)
    processed = supply.get("processed_boxes", 0)
    qr_boxes = supply.get("qr_boxes", [])

    if not qr_boxes:
        await q.edit_message_text(
            "⚠️ <b>Нет добавленных коробок.</b>\n\n"
            "Отсканируйте хотя бы одну коробку перед завершением.",
            parse_mode="HTML",
            reply_markup=InlineKeyboardMarkup([[back_to_main_btn()[0]]]),
        )
        return SUPPLY_WAITING_PHOTO

    total_qr = len(qr_boxes)  # один QR на коробку

    await q.edit_message_text(
        f"⚙️ <b>Генерирую итоговый PDF...</b>\n\n"
        f"📦 Коробок: <b>{processed}</b>\n"
        f"📋 QR-кодов: <b>{total_qr}</b>\n\n"
        "Подождите, это может занять несколько секунд...",
        parse_mode="HTML",
    )

    try:
        pdf_path = generate_qr_pdf(qr_boxes)
        ts_label = datetime.now().strftime("%d.%m.%Y_%H-%M")

        with open(pdf_path, "rb") as fh:
            await context.bot.send_document(
                chat_id=q.message.chat_id,
                document=fh,
                filename=f"Поставка_{ts_label}.pdf",
                caption=(
                    f"📦 <b>Приёмка завершена!</b>\n\n"
                    f"📦 Коробок: <b>{processed}</b>\n"
                    f"📋 QR-кодов в PDF: <b>{total_qr}</b>\n"
                    "<i>Один QR-код = одна коробка (все серийные номера внутри)</i>"
                ),
                parse_mode="HTML",
            )

        # Удаляем сообщение «Генерирую итоговый PDF» — оно больше не нужно
        try:
            await context.bot.delete_message(
                chat_id=q.message.chat_id,
                message_id=q.message.message_id,
            )
        except Exception:
            pass

        try:
            os.remove(pdf_path)
        except Exception:
            pass

    except ImportError as e:
        await context.bot.send_message(
            chat_id=q.message.chat_id,
            text=(
                f"❌ <b>Не установлены библиотеки для генерации QR:</b>\n\n"
                f"<code>{e}</code>\n\n"
                "Установите: <code>pip install qrcode reportlab pillow</code>"
            ),
            parse_mode="HTML",
        )
    except Exception as e:
        logger.error(f"supply_finish PDF error: {e}")
        import traceback; logger.error(traceback.format_exc())
        await context.bot.send_message(
            chat_id=q.message.chat_id,
            text=f"❌ Ошибка при генерации PDF: {e}",
        )

    # Очищаем сессию поставки
    user_sessions.pop(uid, None)

    await context.bot.send_message(
        chat_id=q.message.chat_id,
        text="Выберите следующее действие:",
        reply_markup=main_keyboard(uid),
    )
    return CHOOSING


# ===========================================================================
#  HANDLER — ФОТО СО ШТРИХКОДАМИ
# ===========================================================================

async def handle_barcode_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not await check_access(update):
        return CHOOSING
    uid = update.effective_user.id
    msg = update.message

    photo = msg.photo[-1] if msg.photo else None
    doc = (
        msg.document
        if msg.document and msg.document.mime_type and msg.document.mime_type.startswith("image/")
        else None
    )

    if not photo and not doc:
        await msg.reply_text("❌ Отправьте фотографию или изображение как документ.")
        return WAITING_BARCODE_PHOTO

    status = await msg.reply_text("🔍 Сканирую штрихкоды...")

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    img_path = str(TEMP_DIR / f"scan_{uid}_{ts}.jpg")

    try:
        tg_file = await context.bot.get_file(doc.file_id if doc else photo.file_id)
        await tg_file.download_to_drive(img_path)

        result = scan_barcodes(img_path)

        if not result["ok"]:
            s = check_libs()
            await status.edit_text(
                "❌ <b>Библиотека сканирования не найдена!</b>\n\n"
                + libs_status_text(s)
                + "\n\n💡 После установки перезапустите бота. Проверить: /checklibs",
                parse_mode="HTML",
            )

        elif not result["serials"]:
            total = len(result["all"])
            if total > 0:
                await status.edit_text(
                    f"⚠️ <b>Считано штрихкодов: {total}</b>\n\n"
                    "Все они нечисловые (MAC-адреса и т.п.).\n"
                    "Серийных номеров не обнаружено.\n\n"
                    "Попробуйте отправить фото <b>как документ</b> для лучшего качества.",
                    parse_mode="HTML",
                )
            else:
                await status.edit_text(
                    "⚠️ <b>Штрихкоды не обнаружены.</b>\n\n"
                    "Советы:\n"
                    "• Хорошее освещение, без теней\n"
                    "• Штрихкоды чёткие, не смазанные\n"
                    "• Отправьте фото <b>как документ</b> (без сжатия)\n"
                    "• Подойдите ближе к коробке",
                    parse_mode="HTML",
                )

        else:
            serials = result["serials"]
            skipped = result["skipped"]

            header = (
                f"✅ <b>Серийных номеров: {len(serials)}</b>"
                + (f"  ·  нечисловых пропущено: {skipped}" if skipped else "")
                + "\n\n"
                + "<code>"
                + "\n".join(serials)
                + "</code>"
            )

            if len(header) <= 4096:
                await status.edit_text(header, parse_mode="HTML")
            else:
                await status.edit_text(
                    f"✅ <b>Серийных номеров: {len(serials)}</b>\nОтправляю частями...",
                    parse_mode="HTML"
                )
                chunk: list = []
                part = 1
                for serial in serials:
                    chunk.append(serial)
                    block = "<code>" + "\n".join(chunk) + "</code>"
                    if len(block) > 3800:
                        await context.bot.send_message(
                            chat_id=msg.chat_id,
                            text=f"📋 <b>Часть {part}:</b>\n\n<code>" + "\n".join(chunk[:-1]) + "</code>",
                            parse_mode="HTML",
                        )
                        chunk = [chunk[-1]]
                        part += 1
                if chunk:
                    await context.bot.send_message(
                        chat_id=msg.chat_id,
                        text=f"📋 <b>Часть {part}:</b>\n\n<code>" + "\n".join(chunk) + "</code>",
                        parse_mode="HTML",
                    )

        # Удаляем фото из чата
        try:
            await context.bot.delete_message(chat_id=msg.chat_id, message_id=msg.message_id)
        except Exception as e:
            logger.warning(f"Не удалось удалить фото: {e}")

    except Exception as e:
        logger.error(f"handle_barcode_photo error: {e}")
        import traceback; logger.error(traceback.format_exc())
        await status.edit_text(f"❌ Ошибка: {e}")

    finally:
        if os.path.exists(img_path):
            try: os.remove(img_path)
            except Exception: pass

    await context.bot.send_message(
        chat_id=msg.chat_id,
        text="Выберите следующее действие:",
        reply_markup=scan_after_keyboard()
    )
    return CHOOSING


# ===========================================================================
#  ERROR HANDLER
# ===========================================================================

async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.error(f"Update error: {context.error}")
    if update and update.effective_message:
        await update.effective_message.reply_text("⚠️ Произошла ошибка. Попробуйте ещё раз или /start")


# ===========================================================================
#  FASTAPI — REST API для Mini App
# ===========================================================================

if FASTAPI_AVAILABLE:
    api_app = FastAPI(title="CEK Bot API", docs_url=None, redoc_url=None)
    api_app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    # Сессии для многошагового процесса (гарантия, поставка)
    _api_sessions: dict = {}

    def _api_require_access(user_id: int):
        if not is_allowed(user_id):
            raise HTTPException(status_code=403, detail="Нет доступа. Обратитесь к администратору.")

    def _cleanup_files(*paths):
        for p in paths:
            try:
                if p and os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass

    # ── Ping ────────────────────────────────────────────────────────────────
    @api_app.get("/api/ping")
    async def api_ping():
        return {"status": "ok"}

    # ── Проверка доступа ─────────────────────────────────────────────────────
    @api_app.post("/api/check")
    async def api_check(user_id: int = Form(...)):
        allowed = is_allowed(user_id)
        return {"allowed": allowed, "is_admin": (user_id == ADMIN_ID)}

    # ── Обработка пломб ──────────────────────────────────────────────────────
    @api_app.post("/api/plomb")
    async def api_plomb(user_id: int = Form(...), file: UploadFile = File(...)):
        _api_require_access(user_id)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        suffix = Path(file.filename).suffix.lower() or ".xlsx"
        in_path = str(TEMP_DIR / f"api_plomb_{user_id}_{ts}{suffix}")
        out_path = None
        try:
            with open(in_path, "wb") as fh:
                fh.write(await file.read())
            res = plomb_proc.process_file(in_path)
            if not res["success"]:
                raise HTTPException(status_code=400, detail=res["error"])
            out_path = res["output_path"]
            fname = f"Пломбы_{datetime.now().strftime('%d.%m.%Y')}.xlsx"
            return FileResponse(
                out_path,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                filename=fname,
                background=BackgroundTask(_cleanup_files, in_path, out_path),
            )
        except HTTPException:
            _cleanup_files(in_path, out_path)
            raise
        except Exception as e:
            _cleanup_files(in_path, out_path)
            raise HTTPException(status_code=500, detail=str(e))

    # ── Сканирование штрихкодов ───────────────────────────────────────────────
    @api_app.post("/api/scan")
    async def api_scan(user_id: int = Form(...), file: UploadFile = File(...)):
        _api_require_access(user_id)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        suffix = Path(file.filename).suffix.lower() or ".jpg"
        img_path = str(TEMP_DIR / f"api_scan_{user_id}_{ts}{suffix}")
        try:
            with open(img_path, "wb") as fh:
                fh.write(await file.read())
            result = scan_barcodes(img_path)
            if not result["ok"] and result.get("error") == "no_libs":
                raise HTTPException(status_code=500, detail="Библиотеки сканирования не установлены")
            return {
                "ok": result["ok"],
                "serials": result.get("serials", []),
                "all": result.get("all", []),
                "skipped": result.get("skipped", 0),
                "method": result.get("method", ""),
            }
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
        finally:
            _cleanup_files(img_path)

    # ── Гарантия: шаг 1 — загрузить Excel ────────────────────────────────────
    @api_app.post("/api/guarantee/upload")
    async def api_guarantee_upload(user_id: int = Form(...), file: UploadFile = File(...)):
        _api_require_access(user_id)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        suffix = Path(file.filename).suffix.lower() or ".xlsx"
        path = str(TEMP_DIR / f"api_guar_{user_id}_{ts}{suffix}")
        with open(path, "wb") as fh:
            fh.write(await file.read())
        session_id = str(uuid.uuid4())
        _api_sessions[session_id] = {
            "user_id": user_id,
            "file_path": path,
            "validation": None,
            "all_devices": [],
            "created": datetime.now().isoformat(),
        }
        return {"session_id": session_id}

    # ── Гарантия: шаг 2 — проверить серийные номера ───────────────────────────
    @api_app.post("/api/guarantee/check")
    async def api_guarantee_check(
        user_id: int = Form(...),
        session_id: str = Form(...),
        serials_text: str = Form(...),
    ):
        _api_require_access(user_id)
        sess = _api_sessions.get(session_id)
        if not sess or sess["user_id"] != user_id:
            raise HTTPException(status_code=404, detail="Сессия не найдена")
        try:
            devices = []
            for line in serials_text.strip().splitlines():
                line = line.strip()
                if not line:
                    continue
                if "-" in line:
                    sn, yr = line.split("-", 1)
                else:
                    sn, yr = line, ""
                devices.append({"serial": sn.strip(), "year": yr.strip()})
            if not devices:
                raise HTTPException(status_code=400, detail="Не найдено серийных номеров")
            df = guarantee_proc.load(sess["file_path"])
            val = guarantee_proc.find_devices(df, devices)
            sess["validation"] = val
            sess["all_devices"] = devices
            return {
                "valid": val["valid"],
                "invalid": val["invalid"],
                "total": len(devices),
            }
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    # ── Гарантия: шаг 3 — сгенерировать документы ────────────────────────────
    @api_app.post("/api/guarantee/generate")
    async def api_guarantee_generate(
        user_id: int = Form(...),
        session_id: str = Form(...),
        process_all: str = Form(default="false"),
    ):
        _api_require_access(user_id)
        sess = _api_sessions.get(session_id)
        if not sess or sess["user_id"] != user_id:
            raise HTTPException(status_code=404, detail="Сессия не найдена")
        if not sess.get("validation"):
            raise HTTPException(status_code=400, detail="Сначала выполните проверку серийных номеров")
        try:
            val = sess["validation"]
            use_all = process_all.lower() in ("true", "1", "yes")
            if use_all:
                to_proc = val["valid"] + [
                    {"serial_number": d["serial"], "year": d["year"],
                     "model": "НЕ ПРОШЕЛ ПРОВЕРКУ", "address_full": "", "request_id": ""}
                    for d in val["invalid"]
                ]
            else:
                to_proc = val["valid"]
            if not to_proc:
                raise HTTPException(status_code=400, detail="Нет приборов для обработки")

            tpl = get_template_path(user_id)
            if not tpl.exists():
                raise HTTPException(
                    status_code=400,
                    detail="Шаблон акта не найден. Сначала загрузите шаблон через раздел «Загрузить шаблон»."
                )

            reg_tpl = TEMPLATES_DIR / "реестр_гарантии_шаблон.xlsx"
            if not reg_tpl.exists():
                pd.DataFrame(columns=list("ABCDEFGHIJKLM")).to_excel(reg_tpl, index=False)

            reg_path = guarantee_proc.build_registry(str(reg_tpl), to_proc)
            acts_path = guarantee_proc.build_acts(str(tpl), to_proc)

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            zip_path = str(TEMP_DIR / f"api_docs_{user_id}_{ts}.zip")
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                zf.write(reg_path, f"реестр_{ts}.xlsx")
                zf.write(acts_path, f"акты_{len(to_proc)}шт_{ts}.docx")

            # Чистим сессию и временные файлы
            file_to_clean = sess.get("file_path", "")
            _api_sessions.pop(session_id, None)

            return FileResponse(
                zip_path,
                media_type="application/zip",
                filename=f"Документы_{ts}.zip",
                background=BackgroundTask(_cleanup_files, file_to_clean, reg_path, acts_path, zip_path),
            )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    # ── Загрузка шаблона акта ─────────────────────────────────────────────────
    @api_app.post("/api/template/upload")
    async def api_template_upload(user_id: int = Form(...), file: UploadFile = File(...)):
        _api_require_access(user_id)
        if not file.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="Нужен файл .docx")
        tpl_path = get_template_path(user_id)
        try:
            content = await file.read()
            with open(str(tpl_path), "wb") as fh:
                fh.write(content)
            if not DEFAULT_TEMPLATE_PATH.exists():
                shutil.copy2(str(tpl_path), str(DEFAULT_TEMPLATE_PATH))
            d = Document(str(tpl_path))
            full_text = " ".join(p.text for p in d.paragraphs)
            missing = [ph for ph in ["[[Наименование]]", "[[Номер]]", "[[Год]]", "[[Адрес]]"]
                       if ph not in full_text]
            return {"success": True, "filename": file.filename, "missing_placeholders": missing}
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    # ── Генерация QR-кодов ────────────────────────────────────────────────────
    @api_app.post("/api/qr/generate")
    async def api_qr_generate(
        user_id: int = Form(...),
        boxes_json: str = Form(...),
    ):
        _api_require_access(user_id)
        try:
            boxes = json.loads(boxes_json)
            if not boxes:
                raise HTTPException(status_code=400, detail="Нет данных для генерации")
            pdf_path = generate_qr_pdf(boxes)
            ts = datetime.now().strftime("%d.%m.%Y")
            return FileResponse(
                pdf_path,
                media_type="application/pdf",
                filename=f"QR_коды_{ts}.pdf",
                background=BackgroundTask(_cleanup_files, pdf_path),
            )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    # ── Поставка: сканировать одно фото ──────────────────────────────────────
    @api_app.post("/api/supply/scan")
    async def api_supply_scan(user_id: int = Form(...), file: UploadFile = File(...)):
        _api_require_access(user_id)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        suffix = Path(file.filename).suffix.lower() or ".jpg"
        img_path = str(TEMP_DIR / f"api_supply_{user_id}_{ts}{suffix}")
        try:
            with open(img_path, "wb") as fh:
                fh.write(await file.read())
            result = scan_barcodes(img_path)
            if not result["ok"] and result.get("error") == "no_libs":
                raise HTTPException(status_code=500, detail="Библиотеки сканирования не установлены")
            return {
                "ok": result["ok"],
                "serials": result.get("serials", []),
                "skipped": result.get("skipped", 0),
            }
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))
        finally:
            _cleanup_files(img_path)

    # ── Поставка: сгенерировать QR PDF из коробок ────────────────────────────
    @api_app.post("/api/supply/generate")
    async def api_supply_generate(user_id: int = Form(...), boxes_json: str = Form(...)):
        _api_require_access(user_id)
        try:
            boxes = json.loads(boxes_json)
            if not boxes:
                raise HTTPException(status_code=400, detail="Нет коробок")
            pdf_path = generate_qr_pdf(boxes)
            ts = datetime.now().strftime("%d.%m.%Y")
            return FileResponse(
                pdf_path,
                media_type="application/pdf",
                filename=f"Поставка_{ts}.pdf",
                background=BackgroundTask(_cleanup_files, pdf_path),
            )
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=str(e))

    # ── Список пользователей (только для админа) ─────────────────────────────
    @api_app.get("/api/admin/users")
    async def api_admin_users(user_id: int):
        if user_id != ADMIN_ID:
            raise HTTPException(status_code=403, detail="Только для администратора")
        users = get_all_users()
        return {"users": [
            {"id": uid, "name": info.get("name", ""), "added": info.get("added", "")}
            for uid, info in users.items()
        ]}

    @api_app.post("/api/admin/add")
    async def api_admin_add_user(admin_id: int = Form(...), new_user_id: int = Form(...), name: str = Form(default="")):
        if admin_id != ADMIN_ID:
            raise HTTPException(status_code=403, detail="Только для администратора")
        add_user(new_user_id, name)
        return {"success": True}

    @api_app.post("/api/admin/remove")
    async def api_admin_remove_user(admin_id: int = Form(...), target_user_id: int = Form(...)):
        if admin_id != ADMIN_ID:
            raise HTTPException(status_code=403, detail="Только для администратора")
        remove_user(target_user_id)
        return {"success": True}


# ===========================================================================
#  MAIN
# ===========================================================================

def main():
    print("🤖 Бот запускается...")
    print(f"   Временные файлы: {TEMP_DIR}")
    print(f"   Шаблоны: {TEMPLATES_DIR}")
    print(f"   Файл пользователей: {USERS_FILE}")
    print(f"   Администратор: {ADMIN_ID}")

    s = check_libs()
    print(f"🐍 Python: {s['python']}")
    print(f"   Pillow:   {'✅' if s['pillow'] else '❌  ' + s['pillow_err']}")
    print(f"   pyzbar:   {'✅' if s['pyzbar'] else '❌  ' + s['pyzbar_err']}")
    print(f"   zxingcpp: {'✅' if s['zxingcpp'] else '❌  ' + s['zxingcpp_err']}")
    if s["any"]:
        print(f"✅ Сканирование доступно ({'zxingcpp' if s['zxingcpp'] else 'pyzbar'})")
    else:
        print("⚠️  Сканирование недоступно!")
        print(f"   Установите: {s['python']} -m pip install zxing-cpp pillow")

    if FASTAPI_AVAILABLE:
        print(f"🌐 API для Mini App будет доступен на порту {API_PORT}")
    else:
        print("⚠️  FastAPI не установлен — API для Mini App недоступен.")
        print("   Установите: pip install fastapi uvicorn python-multipart")

    async def _run():
        conv = ConversationHandler(
            entry_points=[
                CommandHandler("start", cmd_start),
                CommandHandler("app", cmd_app),
                CommandHandler("supply", supply_start),
                MessageHandler(filters.StatusUpdate.WEB_APP_DATA, web_app_data_handler),
            ],
            states={
                CHOOSING: [
                    CallbackQueryHandler(btn_callback),
                    CommandHandler("help", cmd_help),
                    CommandHandler("checklibs", cmd_checklibs),
                    CommandHandler("app", cmd_app),
                    MessageHandler(filters.StatusUpdate.WEB_APP_DATA, web_app_data_handler),
                ],
                WAITING_FILES: [
                    MessageHandler(filters.Document.ALL, handle_file),
                    CallbackQueryHandler(btn_callback, pattern="^main$"),
                    CommandHandler("cancel", cmd_cancel),
                ],
                WAITING_SERIALS: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, handle_serials),
                    CallbackQueryHandler(btn_callback, pattern="^main$"),
                    CommandHandler("cancel", cmd_cancel),
                ],
                CONFIRM_DATA: [
                    CallbackQueryHandler(btn_callback, pattern="^(process_all|process_valid|cancel_op)$"),
                    CallbackQueryHandler(btn_callback, pattern="^main$"),
                    CommandHandler("cancel", cmd_cancel),
                ],
                WAITING_TEMPLATE: [
                    MessageHandler(filters.Document.ALL, handle_template_file),
                    CallbackQueryHandler(btn_callback, pattern="^main$"),
                    CommandHandler("cancel", cmd_cancel),
                ],
                WAITING_BARCODE_PHOTO: [
                    MessageHandler(filters.PHOTO | filters.Document.IMAGE, handle_barcode_photo),
                    CallbackQueryHandler(btn_callback),
                    CommandHandler("checklibs", cmd_checklibs),
                    CommandHandler("cancel", cmd_cancel),
                    CommandHandler("start", cmd_start),
                ],
                WAITING_QR_INPUT: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, handle_qr_input),
                    CallbackQueryHandler(btn_callback, pattern="^(qr_continue|qr_finish|main)$"),
                    CommandHandler("cancel", cmd_cancel),
                ],
                SUPPLY_WAITING_PHOTO: [
                    MessageHandler(filters.PHOTO | filters.Document.IMAGE, handle_supply_photo),
                    CallbackQueryHandler(btn_callback, pattern="^(supply_photo_ok|supply_photo_retry|supply_edit_serials|main)$"),
                    CommandHandler("cancel", cmd_cancel),
                    CommandHandler("start", cmd_start),
                ],
                SUPPLY_WAITING_NUMBER: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, handle_supply_number),
                    CallbackQueryHandler(btn_callback, pattern="^(supply_add_more|supply_finish|main)$"),
                    CommandHandler("cancel", cmd_cancel),
                    CommandHandler("start", cmd_start),
                ],
                SUPPLY_EDITING_SERIALS: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, handle_supply_edit_input),
                    CallbackQueryHandler(btn_callback, pattern="^(supply_edit_serials|main)$"),
                    CommandHandler("cancel", cmd_cancel),
                    CommandHandler("start", cmd_start),
                ],
                ADMIN_PANEL: [
                    CallbackQueryHandler(btn_callback),
                    MessageHandler(filters.TEXT & ~filters.COMMAND, admin_receive_user_id),
                    CommandHandler("cancel", cmd_cancel),
                ],
                ADMIN_ADD_USER: [
                    MessageHandler(filters.TEXT & ~filters.COMMAND, admin_receive_user_id),
                    CallbackQueryHandler(btn_callback),
                    CommandHandler("cancel", cmd_cancel),
                ],
            },
            fallbacks=[
                CommandHandler("cancel", cmd_cancel),
                CommandHandler("start", cmd_start),
            ],
        )

        application = (
            Application.builder()
            .token(BOT_TOKEN)
            .read_timeout(120)
            .write_timeout(180)
            .connect_timeout(60)
            .pool_timeout(60)
            .build()
        )
        application.add_handler(conv)
        application.add_handler(CommandHandler("help", cmd_help))
        application.add_handler(CommandHandler("checklibs", cmd_checklibs))
        application.add_handler(CommandHandler("myid", cmd_myid))
        application.add_handler(CommandHandler("app", cmd_app))
        application.add_handler(MessageHandler(filters.StatusUpdate.WEB_APP_DATA, web_app_data_handler))
        application.add_handler(CommandHandler("supply", supply_start))
        application.add_error_handler(error_handler)

        if FASTAPI_AVAILABLE:
            uvicorn_config = uvicorn.Config(
                api_app, host="0.0.0.0", port=API_PORT,
                log_level="warning", access_log=False,
            )
            uvicorn_server = uvicorn.Server(uvicorn_config)

            async with application:
                await application.start()
                await application.updater.start_polling(allowed_updates=Update.ALL_TYPES)
                print(f"🤖 Бот запущен!")
                print(f"🌐 API запущен: http://0.0.0.0:{API_PORT}  (Mini App может подключаться)")
                print("   Ctrl+C для остановки.")
                await uvicorn_server.serve()
                await application.updater.stop()
                await application.stop()
        else:
            print("🤖 Бот запущен! Ctrl+C для остановки.")
            application.run_polling(allowed_updates=Update.ALL_TYPES)

    asyncio.run(_run())


if __name__ == "__main__":
    main()